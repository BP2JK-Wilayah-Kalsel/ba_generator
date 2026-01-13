import os
import re
import json
from flask import Flask, render_template, request, jsonify, send_file, flash, redirect, url_for, send_from_directory
from docx import Document
from docx.shared import Inches
from datetime import datetime
import calendar
import zipfile
import tempfile
import shutil
from pathlib import Path
import locale
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml import OxmlElement
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from copy import copy

# Import SPSE crawler
from spse_crawler import create_spse_endpoint

# Set locale untuk bahasa Indonesia
try:
    locale.setlocale(locale.LC_TIME, 'id_ID.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_TIME, 'Indonesian_Indonesia.1252')
    except:
        pass

app = Flask(__name__)
app.secret_key = 'ba_generator_secret_2025'

# Global storage untuk hasil processing
PROCESSED_FILES_DIR = os.path.join(os.getcwd(), 'processed_results')
os.makedirs(PROCESSED_FILES_DIR, exist_ok=True)

def clean_old_processed_files():
    """
    Auto-clean file lama di processed_results untuk menghindari penumpukan
    Fungsi ini dipanggil setiap kali akan generate dokumen baru
    """
    try:
        if os.path.exists(PROCESSED_FILES_DIR):
            # Hapus semua file .docx yang lebih dari 1 jam
            import time
            current_time = time.time()
            one_hour = 3600  # 1 jam dalam detik
            
            for filename in os.listdir(PROCESSED_FILES_DIR):
                filepath = os.path.join(PROCESSED_FILES_DIR, filename)
                if os.path.isfile(filepath):
                    # Check jika file lebih tua dari 1 jam
                    file_age = current_time - os.path.getmtime(filepath)
                    if file_age > one_hour:
                        try:
                            os.remove(filepath)
                            print(f"[AUTO-CLEAN] Removed old file: {filename}")
                        except Exception as e:
                            print(f"[AUTO-CLEAN] Failed to remove {filename}: {e}")
        
        print(f"[AUTO-CLEAN] Cleaned old files from processed_results")
    except Exception as e:
        print(f"[AUTO-CLEAN] Error during cleanup: {e}")

# Document type mapping untuk nomor urut surat
DOCUMENT_TYPES = {
    '00': 'Cover',
    '06': 'Berita Acara Pemberian Penjelasan Kualifikasi',
    '10': 'Berita Acara Hasil Evaluasi Kualifikasi',
    '11': 'Berita Acara Penetapan Daftar Pendek',
    '12': 'Pengumuman Daftar Pendek',
    '13': 'Berita Acara Jawab Sanggah Prakualifikasi',
    '14': 'Berita Acara Pemberian Penjelasan Seleksi',
    '17': 'Berita Acara Hasil Evaluasi Administrasi Dan Teknis',
    '19': 'Berita Acara Hasil Evaluasi Biaya',
    '20': 'Berita Acara Kombinasi Teknis Dan Biaya',
    '21': 'Surat Klarifikasi Personel',
    '22': 'Berita Acara Klarifikasi Penetapan Pemenang',
    '22-LHP': 'Berita Acara Hasil Penelitian',
    '24': 'Berita Acara Penetapan Pemenang',
    '25': 'Berita Acara Pengumuman Pemenang',
    '26': 'Berita Acara Jawab Sanggah Seleksi',
    '27': 'Berita Acara Klarifikasi Dan Negosiasi Teknis Dan Biaya',
    '27-2': 'Daftar Hadir Klarifikasi Dan Negosiasi',
    '28': 'Berita Acara Hasil Pemilihan',
    '29': 'Surat Penyampaian BAHP',
    '96': 'Surat pernyataan Klarifikasi personil dan paket 1 dan 2',
    '97': 'Berita Acara Seleksi Gagal',
    '99': 'TTD Pokja',
}

# Indonesian day and month names
DAY_NAMES = ['Minggu', 'Senin', 'Selasa', 'Rabu', 'Kamis', 'Jumat', 'Sabtu']
MONTH_NAMES = [
    'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni',
    'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember'
]

def terbilang(angka, capitalize_each_word=False):
    """Mengubah angka menjadi terbilang dalam bahasa Indonesia"""
    if angka == 0:
        result = "nol"
        return result.title() if capitalize_each_word else result
    
    satuan = ["", "satu", "dua", "tiga", "empat", "lima", "enam", "tujuh", "delapan", "sembilan"]
    belasan = ["sepuluh", "sebelas", "dua belas", "tiga belas", "empat belas", "lima belas", 
               "enam belas", "tujuh belas", "delapan belas", "sembilan belas"]
    puluhan = ["", "", "dua puluh", "tiga puluh", "empat puluh", "lima puluh", 
               "enam puluh", "tujuh puluh", "delapan puluh", "sembilan puluh"]
    
    def konversi_ratusan(n):
        hasil = ""
        if n >= 100:
            if n >= 200:
                hasil += satuan[n // 100] + " ratus "
            else:
                hasil += "seratus "
            n %= 100
        
        if n >= 20:
            hasil += puluhan[n // 10] + " "
            n %= 10
        elif n >= 10:
            hasil += belasan[n - 10] + " "
            n = 0
        
        if n > 0:
            hasil += satuan[n] + " "
        
        return hasil.strip()
    
    def konversi_ribuan(n):
        if n == 0:
            return ""
        elif n == 1:
            return "seribu "
        else:
            return konversi_ratusan(n) + " ribu "
    
    def konversi_jutaan(n):
        if n == 0:
            return ""
        elif n == 1:
            return "satu juta "
        else:
            return konversi_ratusan(n) + " juta "
    
    def konversi_miliaran(n):
        if n == 0:
            return ""
        elif n == 1:
            return "satu miliar "
        else:
            return konversi_ratusan(n) + " miliar "
    
    def konversi_triliunan(n):
        if n == 0:
            return ""
        elif n == 1:
            return "satu triliun "
        else:
            return konversi_ratusan(n) + " triliun "
    
    hasil = ""
    
    # Triliun
    if angka >= 1000000000000:
        triliun = angka // 1000000000000
        hasil += konversi_triliunan(triliun)
        angka %= 1000000000000
    
    # Miliar
    if angka >= 1000000000:
        miliar = angka // 1000000000
        hasil += konversi_miliaran(miliar)
        angka %= 1000000000
    
    # Juta
    if angka >= 1000000:
        juta = angka // 1000000
        hasil += konversi_jutaan(juta)
        angka %= 1000000
    
    # Ribu
    if angka >= 1000:
        ribu = angka // 1000
        hasil += konversi_ribuan(ribu)
        angka %= 1000
    
    # Ratusan, puluhan, satuan
    if angka > 0:
        hasil += konversi_ratusan(angka) + " "
    
    result = hasil.strip()
    return result.title() if capitalize_each_word else result

def format_currency(amount):
    """Format angka menjadi format rupiah dengan ,00 di akhir"""
    if amount == 0:
        return "Rp 0,00"
    
    # Convert to string and add thousand separators with ,00 at the end
    formatted = f"Rp {amount:,.2f}".replace(',', 'TEMP').replace('.', ',').replace('TEMP', '.')
    return formatted

def generate_comprehensive_keywords(form_data):
    """Generate comprehensive keywords from form data"""
    keywords = {}
    
    # Basic information
    keywords['nomor_sk_pokja'] = form_data.get('nomor_sk_pokja', '')
    keywords['tanggal_sk_pokja'] = format_date_indonesian(form_data.get('tanggal_sk_pokja', ''))
    keywords['nomor_sk_timlak'] = form_data.get('nomor_sk_timlak', '')
    keywords['tanggal_sk_timlak'] = format_date_indonesian(form_data.get('tanggal_sk_timlak', ''))
    keywords['kode_pokja'] = form_data.get('kode_pokja', '')
    keywords['tahun_surat'] = form_data.get('tahun_surat', '')
    keywords['tahun_anggaran'] = form_data.get('tahun_anggaran', '')
    keywords['kode_tender'] = form_data.get('kode_tender', '')
    keywords['nama_paket'] = form_data.get('nama_paket', '')
    keywords['klpd'] = form_data.get('klpd', '')
    keywords['unit_organisasi'] = form_data.get('unit_organisasi', '')
    keywords['balai'] = form_data.get('balai', '')
    keywords['satuan_kerja'] = form_data.get('satuan_kerja', '')
    keywords['kegiatan'] = form_data.get('kegiatan', '')
    keywords['jenis_pengadaan'] = form_data.get('jenis_pengadaan', '')
    keywords['metode_pengadaan'] = form_data.get('metode_pengadaan', '')
    keywords['sumber_dana'] = form_data.get('sumber_dana', '')
    
    # Values
    try:
        nilai_pagu = int(form_data.get('nilai_pagu', '0') or '0')
        nilai_hps = int(form_data.get('nilai_hps', '0') or '0')
    except:
        nilai_pagu = nilai_hps = 0
    
    keywords['nilai_pagu'] = format_currency(nilai_pagu)
    # Terbilang untuk nilai pagu: gunakan huruf kecil semua
    keywords['terbilang_pagu'] = terbilang(nilai_pagu, False) + ' Rupiah' if nilai_pagu > 0 else ''
    keywords['nilai_hps'] = format_currency(nilai_hps)
    # Terbilang untuk nilai HPS: gunakan huruf kecil semua
    keywords['terbilang_hps'] = terbilang(nilai_hps, False) + ' Rupiah' if nilai_hps > 0 else ''
    
    # Pokja pemilihan
    keywords['pokja_pemilihan'] = f"Kelompok Kerja Pemilihan {keywords['kode_pokja']} BP2JK Wilayah Kalimantan Selatan Tahun Anggaran {keywords['tahun_anggaran']}"
    
    # POKJA roles - Updated to handle new structure
    # Handle ketua data (could be JSON string from dropdown or direct field)
    ketua_data = form_data.get('ketua_pokja', '')
    if ketua_data:
        try:
            # Try to parse as JSON (new system)
            ketua_obj = json.loads(ketua_data)
            keywords['ketua_pokja'] = ketua_obj.get('nama', '')
            keywords['nip_ketua_pokja'] = ketua_obj.get('nip', '')
            keywords['email_ketua_pokja'] = ketua_obj.get('email', '')
        except (json.JSONDecodeError, TypeError):
            # Fallback to direct field value (old system or manual input)
            keywords['ketua_pokja'] = ketua_data
            keywords['nip_ketua_pokja'] = form_data.get('nip_ketua_pokja', '')
            keywords['email_ketua_pokja'] = form_data.get('email_ketua_pokja', '')
    else:
        keywords['ketua_pokja'] = ''
        keywords['nip_ketua_pokja'] = ''
        keywords['email_ketua_pokja'] = ''
    
    # Handle sekre data (could be JSON string from dropdown or direct field)
    sekre_data = form_data.get('sekre_pokja', '')
    if sekre_data:
        try:
            # Try to parse as JSON (new system)
            sekre_obj = json.loads(sekre_data)
            keywords['sekre_pokja'] = sekre_obj.get('nama', '')
            keywords['nip_sekre_pokja'] = sekre_obj.get('nip', '')
            keywords['email_sekre_pokja'] = sekre_obj.get('email', '')
        except (json.JSONDecodeError, TypeError):
            # Fallback to direct field value (old system or manual input)
            keywords['sekre_pokja'] = sekre_data
            keywords['nip_sekre_pokja'] = form_data.get('nip_sekre_pokja', '')
            keywords['email_sekre_pokja'] = form_data.get('email_sekre_pokja', '')
    else:
        keywords['sekre_pokja'] = ''
        keywords['nip_sekre_pokja'] = ''
        keywords['email_sekre_pokja'] = ''
    
    # Handle anggota from multiple select
    # Initialize all anggota slots
    for i in range(3, 8):  # anggota3 to anggota7
        keywords[f'anggota{i}_pokja'] = ''
        keywords[f'nip_anggota{i}_pokja'] = ''
        keywords[f'email_anggota{i}_pokja'] = ''
    
    # Process anggota_pokja multiple selections
    anggota_selections = form_data.getlist('anggota_pokja') if hasattr(form_data, 'getlist') else (
        form_data.get('anggota_pokja', []) if isinstance(form_data.get('anggota_pokja', []), list) 
        else [form_data.get('anggota_pokja', '')] if form_data.get('anggota_pokja', '') else []
    )
    
    anggota_index = 3  # Start from anggota3
    for anggota_data in anggota_selections:
        if anggota_index > 7:  # Maximum 5 additional anggota (3-7)
            break
        if anggota_data:
            try:
                # Try to parse as JSON (new system)
                anggota_obj = json.loads(anggota_data)
                keywords[f'anggota{anggota_index}_pokja'] = anggota_obj.get('nama', '')
                keywords[f'nip_anggota{anggota_index}_pokja'] = anggota_obj.get('nip', '')
                keywords[f'email_anggota{anggota_index}_pokja'] = anggota_obj.get('email', '')
                anggota_index += 1
            except (json.JSONDecodeError, TypeError):
                # Fallback to direct field value (old system)
                keywords[f'anggota{anggota_index}_pokja'] = anggota_data
                keywords[f'nip_anggota{anggota_index}_pokja'] = form_data.get(f'nip_anggota{anggota_index}_pokja', '')
                keywords[f'email_anggota{anggota_index}_pokja'] = form_data.get(f'email_anggota{anggota_index}_pokja', '')
                anggota_index += 1
    
    # Also check for direct field inputs (backward compatibility)
    for i in range(3, 8):
        if not keywords[f'anggota{i}_pokja']:  # Only if not already filled from selections
            keywords[f'anggota{i}_pokja'] = form_data.get(f'anggota{i}_pokja', '')
            keywords[f'nip_anggota{i}_pokja'] = form_data.get(f'nip_anggota{i}_pokja', '')
            keywords[f'email_anggota{i}_pokja'] = form_data.get(f'email_anggota{i}_pokja', '')
    
    # Generate document-specific date keywords (turunan dari format_tanggal_XX)
    # List of document numbers that might have dates
    doc_numbers = ['00', '06', '10', '11', '12', '13', '14', '17', '19', '20', '21', '22', '22-LHP', '24', '25', '26', '27', '27-2', '28', '29', '96', '97', '99']
    
    for doc_num in doc_numbers:
        # Check if format_tanggal_XX exists in form
        date_key = f'format_tanggal_{doc_num}'
        date_value = form_data.get(date_key, '')
        
        if date_value:
            # Generate all derivative date keywords
            date_keywords = generate_date_keywords(date_value, doc_num)
            keywords.update(date_keywords)

    return keywords


def generate_date_keywords(date_string, doc_num):
    """
    Generate all derivative date keywords from format_tanggal_XX
    
    Args:
        date_string: Date in format 'YYYY-MM-DD' or 'DD/MM/YYYY'
        doc_num: Document number (e.g., '28', '24', '22-LHP')
    
    Returns:
        dict: Dictionary with all date-related keywords
    """
    keywords = {}
    
    if not date_string:
        return keywords
    
    try:
        # Parse date - handle both YYYY-MM-DD and DD/MM/YYYY formats
        if '/' in date_string:
            # DD/MM/YYYY format
            day, month, year = date_string.split('/')
            date_obj = datetime(int(year), int(month), int(day))
        else:
            # YYYY-MM-DD format
            date_obj = datetime.strptime(date_string, '%Y-%m-%d')
        
        # Generate all derivative keywords
        # {format_tanggal_XX} - keep original
        keywords[f'format_tanggal_{doc_num}'] = date_string
        
        # {tanggal_bulan_tahun_XX} - "13 Agustus 2025"
        keywords[f'tanggal_bulan_tahun_{doc_num}'] = f"{date_obj.day} {MONTH_NAMES[date_obj.month - 1]} {date_obj.year}"
        
        # {hari_surat_XX} - "Rabu"
        keywords[f'hari_surat_{doc_num}'] = DAY_NAMES[date_obj.weekday()]
        
        # {tanggal_sebut_XX} - "Tiga Belas"
        keywords[f'tanggal_sebut_{doc_num}'] = terbilang(date_obj.day, True)
        
        # {bulan_sebut_XX} - "Agustus"
        keywords[f'bulan_sebut_{doc_num}'] = MONTH_NAMES[date_obj.month - 1]
        
        # {tahun_sebut_XX} - "Dua Ribu Dua Puluh Lima"
        keywords[f'tahun_sebut_{doc_num}'] = terbilang(date_obj.year, True)
        
    except (ValueError, IndexError) as e:
        # If date parsing fails, return empty keywords
        pass
    
    return keywords

def format_date_indonesian(date_string):
    """Format date to Indonesian format"""
    if not date_string:
        return ''
    try:
        date_obj = datetime.strptime(date_string, '%Y-%m-%d')
        day = date_obj.day
        month = MONTH_NAMES[date_obj.month - 1]
        year = date_obj.year
        return f"{day} {month} {year}"
    except:
        return date_string

def generate_keywords_from_form(form_data):
    """Generate all keywords from form data (legacy function)"""
    keywords = {}
    
    # Basic form fields
    keywords['kode_pokja'] = form_data.get('kode_pokja', '')
    keywords['tahun'] = form_data.get('tahun', '')
    keywords['nama_paket'] = form_data.get('nama_paket', '')
    keywords['klpd'] = form_data.get('klpd', '')
    keywords['satuan_kerja'] = form_data.get('satuan_kerja', '')
    
    # Generate nomor surat
    doc_type = form_data.get('document_type', '')
    keywords['nomor_surat'] = f"PB0301-Bp2jk17/POKJA-{keywords['kode_pokja']}/{keywords['tahun']}/{doc_type}"
    
    # Custom variables
    # Collect custom variables from form
    for key, value in form_data.items():
        if key.startswith('custom_var_name_'):
            var_index = key.replace('custom_var_name_', '')
            var_value_key = f'custom_var_value_{var_index}'
            var_name = value
            var_value = form_data.get(var_value_key, '')
            
            if var_name and var_value:
                keywords[var_name] = var_value
    
    return keywords

def replace_keywords_in_text(text, keywords, keyword_details=None):
    """Helper function to perform keyword replacement on a string."""
    if keyword_details is None:
        keyword_details = {}

    # Build a regex pattern to find all keywords: {key1}|{key2}|...
    # This is more efficient than looping and replacing one by one.
    if not keywords:
        return text, 0

    # Escape special regex characters in keys and create the pattern
    # The pattern looks for {key}
    pattern = re.compile(r'\{(' + '|'.join(re.escape(key) for key in keywords.keys()) + r')\}')

    replacements_count = 0

    def replacer(match):
        nonlocal replacements_count
        key = match.group(1) # The key without braces
        if key in keywords:
            value = str(keywords[key])
            # Update keyword details for reporting
            if key in keyword_details:
                keyword_details[key]['count'] += 1
            else:
                keyword_details[key] = {'value': value, 'count': 1}
            replacements_count += 1
            return value
        return match.group(0) # Return original if key not found (should not happen with this pattern)

    # Perform all replacements in one pass
    new_text = pattern.sub(replacer, text)
    
    return new_text, replacements_count


def process_hyperlinks_in_paragraph(paragraph, keywords, keyword_details=None):
    """
    Process placeholders inside hyperlinks in a paragraph.
    Hyperlinks are not included in paragraph.runs, so they need special handling.
    Returns the number of replacements made.
    """
    if keyword_details is None:
        keyword_details = {}
    
    total_replacements = 0
    
    # Get paragraph XML element
    para_xml = paragraph._element
    
    # Find all hyperlinks in the paragraph
    hyperlinks = para_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
    
    if not hyperlinks:
        return 0
    
    # Pattern to find placeholders
    pattern = re.compile(r'\{([a-zA-Z0-9_\-]+)\}')
    

    for hyperlink in hyperlinks:
        # Get all text elements in the hyperlink
        text_elements = hyperlink.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')

        # Combine text from all elements
        full_text = "".join([t.text for t in text_elements if t.text])

        # Find placeholders
        matches = list(pattern.finditer(full_text))

        if not matches:
            continue

        # Process each match
        for match in matches:
            keyword = match.group(1)

            # Case-insensitive lookup
            matched_key = None
            for k in keywords.keys():
                if k.lower() == keyword.lower():
                    matched_key = k
                    break

            if matched_key:
                value_to_replace = str(keywords[matched_key])

                # Replace in the first text element (simple approach)
                if text_elements:
                    text_elements[0].text = value_to_replace

                    # Clear other text elements in the hyperlink
                    for t in text_elements[1:]:
                        t.text = ""

                    # Update hyperlink address (mailto or URL)
                    # Get the relationship ID from hyperlink
                    r_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    
                    if r_id:
                        try:
                            # Get document part relationships
                            doc_part = paragraph.part
                            
                            # Access the relationship via rels
                            if hasattr(doc_part, 'rels') and r_id in doc_part.rels:
                                rel = doc_part.rels[r_id]
                                
                                # Check if it has target_ref (read-only property)
                                if hasattr(rel, 'target_ref'):
                                    current_target = rel.target_ref
                                    
                                    if current_target:
                                        # If it's a mailto link
                                        if current_target.startswith('mailto:'):
                                            # Update mailto address via _target attribute
                                            new_target = f"mailto:{value_to_replace}"
                                            rel._target = new_target
                                            print(f"[HYPERLINK] Updated mailto: {current_target} -> {new_target}")
                                        # If URL contains placeholders
                                        elif '{' in current_target:
                                            # Replace all placeholders in URL
                                            updated_url = current_target
                                            for kw, val in keywords.items():
                                                updated_url = updated_url.replace(f'{{{kw}}}', str(val))
                                            rel._target = updated_url
                                            print(f"[HYPERLINK] Updated URL: {current_target} -> {updated_url}")
                                    
                        except Exception as e:
                            print(f"[HYPERLINK] Warning: Could not update link address for {matched_key}: {e}")

                    # Update keyword details
                    if matched_key in keyword_details:
                        keyword_details[matched_key]['count'] += 1
                    else:
                        keyword_details[matched_key] = {'value': value_to_replace, 'count': 1}

                    total_replacements += 1

    return total_replacements


def process_paragraph_keywords(paragraph, keywords, keyword_details=None):
    """
    Processes keywords in a paragraph, handling placeholders that might be split across multiple 'runs'.
    This is a robust method that preserves formatting and works for headers, footers, and body text.
    Also processes hyperlinks which are not included in paragraph.runs.
    """
    if keyword_details is None:
        keyword_details = {}
    
    # Process hyperlinks first (they're not in runs)
    hyperlink_replacements = process_hyperlinks_in_paragraph(paragraph, keywords, keyword_details)
    
    if not paragraph.runs:
        return hyperlink_replacements
    
    # --- Phase 1: Combine runs to find full keywords ---
    full_text = "".join(run.text for run in paragraph.runs)
    
    # If no '{' is in the text, we can skip this paragraph entirely.
    if '{' not in full_text:
        return 0

    # Find all occurrences of {keyword} in the combined text
    # Updated regex to be more flexible: allows letters, numbers, underscores, and hyphens
    pattern = re.compile(r'\{([a-zA-Z0-9_\-]+)\}')
    matches = list(pattern.finditer(full_text))
    
    if not matches:
        return 0

    total_replacements_in_paragraph = 0
    
    # --- Phase 2: Apply replacements ---
    # We iterate through matches in reverse to avoid messing up indices of later matches.
    for match in reversed(matches):
        keyword = match.group(1)
        
        # Case-insensitive keyword lookup
        matched_key = None
        for k in keywords.keys():
            if k.lower() == keyword.lower():
                matched_key = k
                break
        
        if matched_key:
            value_to_replace = str(keywords[matched_key])
            start, end = match.span()

            # Find which runs are affected by this match
            affected_runs = []
            current_pos = 0
            for i, run in enumerate(paragraph.runs):
                run_len = len(run.text)
                # Check if the run overlaps with the match
                if current_pos < end and current_pos + run_len > start:
                    affected_runs.append((i, run))
                current_pos += run_len
            
            if not affected_runs:
                continue

            # Replace the content
            # The first affected run gets the new value and keeps its formatting.
            # Subsequent affected runs are cleared.
            first_run_index, first_run = affected_runs[0]
            
            # Update keyword details for reporting
            if matched_key in keyword_details:
                keyword_details[matched_key]['count'] += 1
            else:
                keyword_details[matched_key] = {'value': value_to_replace, 'count': 1}
            
            total_replacements_in_paragraph += 1

            # To preserve formatting, we need to handle the text within the runs carefully.
            # Find where the match starts relative to the beginning of the paragraph
            run_start_pos = 0
            first_run_original_text = ""
            for i, run in enumerate(paragraph.runs):
                if i == first_run_index:
                    first_run_original_text = run.text
                    break
                run_start_pos += len(run.text)

            # The text in the first run is composed of:
            # 1. The part of the run before the placeholder started.
            # 2. The replacement value.
            # 3. The part of the run after the placeholder ended (might come from last run).
            
            # Calculate the part of the first run that is *before* the match
            prefix_len = start - run_start_pos
            prefix = first_run_original_text[:prefix_len] if prefix_len > 0 else ""

            # Calculate the part of the last affected run that is *after* the match
            last_run_index, last_run = affected_runs[-1]
            run_end_pos = 0
            for i, run in enumerate(paragraph.runs):
                run_end_pos += len(run.text)
                if i == last_run_index:
                    break
            
            suffix_len = run_end_pos - end
            suffix = last_run.text[-suffix_len:] if suffix_len > 0 else ""

            # Set the text for the first run
            first_run.text = prefix + value_to_replace + suffix

            # Clear text from other affected runs
            for i, run in affected_runs[1:]:
                run.text = ""
            
            # Rebuild full_text after this replacement for subsequent matches
            full_text = "".join(run.text for run in paragraph.runs)

    return total_replacements_in_paragraph


def process_docx_comprehensive(file_path, keywords, output_path, deleted_documents=None, keywords_to_delete_rows=None):
    """Process DOCX file with comprehensive keyword replacement and row deletion support"""
    try:
        doc = Document(file_path)
        total_replacements = 0
        log_entries = []
        keyword_details = {}  # Track details for each keyword
        
        # Ensure keywords_to_delete_rows is a list
        if keywords_to_delete_rows is None:
            keywords_to_delete_rows = []
        
        # Process main document paragraphs
        for paragraph in doc.paragraphs:
            before_text = paragraph.text
            replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
            after_text = paragraph.text
            
            if replacements > 0:
                total_replacements += replacements
                log_entries.append({
                    'context': 'paragraph',
                    'before': before_text,
                    'after': after_text,
                    'replacements': replacements
                })
        
        # Process tables with deletion support
        for table in doc.tables:
            rows_to_delete = []
            
            for row_idx, row in enumerate(table.rows):
                row_should_delete = False
                
                # Check each cell for "delete" keyword or deleted document references
                for cell in row.cells:
                    cell_text = cell.text.lower().strip()
                    
                    # Check if cell contains "delete"
                    if cell_text == 'delete':
                        row_should_delete = True
                        break
                    
                    # Check if cell contains reference to deleted documents
                    if deleted_documents:
                        for deleted_doc in deleted_documents:
                            if deleted_doc in cell_text:
                                row_should_delete = True
                                break
                    
                    # NEW: Check if cell contains keywords that should trigger row deletion
                    if keywords_to_delete_rows:
                        for keyword_to_delete in keywords_to_delete_rows:
                            # Check if placeholder exists in this cell
                            if f"{{{keyword_to_delete}}}" in cell.text:
                                row_should_delete = True
                                break
                    
                    if row_should_delete:
                        break
                    
                    # Process keywords in cell paragraphs
                    for paragraph in cell.paragraphs:
                        before_text = paragraph.text
                        replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                        after_text = paragraph.text
                        
                        if replacements > 0:
                            total_replacements += replacements
                            log_entries.append({
                                'context': 'table_cell',
                                'before': before_text,
                                'after': after_text,
                                'replacements': replacements
                            })
                
                if row_should_delete:
                    rows_to_delete.append(row_idx)
            
            # Delete rows in reverse order to maintain correct indices
            for row_idx in reversed(rows_to_delete):
                try:
                    table._element.remove(table.rows[row_idx]._element)
                    log_entries.append({
                        'context': 'table_row_deleted',
                        'before': f'Row {row_idx + 1} in table',
                        'after': 'DELETED',
                        'replacements': 0
                    })
                except:
                    pass
        
        # Process headers and footers
        for section in doc.sections:
            # Headers
            if section.header:
                for paragraph in section.header.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({
                            'context': 'header',
                            'before': before_text,
                            'after': after_text,
                            'replacements': replacements
                        })
                # Header tables
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                before_text = paragraph.text
                                replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                                after_text = paragraph.text
                                if replacements > 0:
                                    total_replacements += replacements
                                    log_entries.append({
                                        'context': 'header_table_cell',
                                        'before': before_text,
                                        'after': after_text,
                                        'replacements': replacements
                                    })
            
            # First Page Header
            if section.first_page_header:
                for paragraph in section.first_page_header.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({
                            'context': 'first_page_header',
                            'before': before_text,
                            'after': after_text,
                            'replacements': replacements
                        })
                # First page header tables
                for table in section.first_page_header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                before_text = paragraph.text
                                replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                                after_text = paragraph.text
                                if replacements > 0:
                                    total_replacements += replacements
                                    log_entries.append({
                                        'context': 'first_page_header_table_cell',
                                        'before': before_text,
                                        'after': after_text,
                                        'replacements': replacements
                                    })
            
            # Even Page Header
            if section.even_page_header:
                for paragraph in section.even_page_header.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({
                            'context': 'even_page_header',
                            'before': before_text,
                            'after': after_text,
                            'replacements': replacements
                        })
                # Even page header tables
                for table in section.even_page_header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                before_text = paragraph.text
                                replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                                after_text = paragraph.text
                                if replacements > 0:
                                    total_replacements += replacements
                                    log_entries.append({
                                        'context': 'even_page_header_table_cell',
                                        'before': before_text,
                                        'after': after_text,
                                        'replacements': replacements
                                    })
            
            # Footers
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({
                            'context': 'footer',
                            'before': before_text,
                            'after': after_text,
                            'replacements': replacements
                        })
                # Footer tables
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                before_text = paragraph.text
                                replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                                after_text = paragraph.text
                                if replacements > 0:
                                    total_replacements += replacements
                                    log_entries.append({
                                        'context': 'footer_table_cell',
                                        'before': before_text,
                                        'after': after_text,
                                        'replacements': replacements
                                    })
        
        # Save the processed document
        doc.save(output_path)
        
        return True, {
            'total_replacements': total_replacements,
            'log_entries': log_entries,
            'keyword_details': keyword_details
        }
        
    except Exception as e:
        return False, {'error': str(e)}

def process_docx_keywords(file_path, keywords, output_path):
    """Process DOCX file and replace all keywords (legacy function)"""
    try:
        doc = Document(file_path)
        total_replacements = 0
        log_entries = []
        keyword_details = {}  # Track details for each keyword
        
        # Process main document paragraphs
        for paragraph in doc.paragraphs:
            before_text = paragraph.text
            replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
            after_text = paragraph.text
            
            if replacements > 0:
                total_replacements += replacements
                log_entries.append({
                    'context': 'paragraph',
                    'before': before_text,
                    'after': after_text,
                    'replacements': replacements
                })
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        before_text = paragraph.text
                        replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                        after_text = paragraph.text
                        
                        if replacements > 0:
                            total_replacements += replacements
                            log_entries.append({
                                'context': 'table_cell',
                                'before': before_text,
                                'after': after_text,
                                'replacements': replacements
                            })
        
        # Process headers and footers
        for section in doc.sections:
            # Headers - Process ALL header types (default, first_page, even_page)
            if section.header:
                for paragraph in section.header.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({
                            'context': 'header',
                            'before': before_text,
                            'after': after_text,
                            'replacements': replacements
                        })
            
            # First Page Header (different header for first page)
            if section.first_page_header:
                for paragraph in section.first_page_header.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({
                            'context': 'first_page_header',
                            'before': before_text,
                            'after': after_text,
                            'replacements': replacements
                        })
            
            # Even Page Header (different header for even pages)
            if section.even_page_header:
                for paragraph in section.even_page_header.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({
                            'context': 'even_page_header',
                            'before': before_text,
                            'after': after_text,
                            'replacements': replacements
                        })
            
            # Footers
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({
                            'context': 'footer',
                            'before': before_text,
                            'after': after_text,
                            'replacements': replacements
                        })
        
        # Save the processed document
        doc.save(output_path)
        
        return True, {
            'total_replacements': total_replacements,
            'log_entries': log_entries,
            'keyword_details': keyword_details
        }
        
    except Exception as e:
        return False, {'error': str(e)}

@app.route('/')
def index():
    """Home page with menu navigation"""
    return render_template('home.html')

@app.route('/ba-pokja-konsultan')
def ba_pokja_konsultan():
    """BA Konsultan page - main application"""
    return render_template('ba_pokja_konsultan.html')

@app.route('/ba-timlak-konsultan')
def ba_timlak_konsultan():
    """BA Tim Pelaksana (Konsultan) page - main application"""
    return render_template('ba_timlak_konsultan.html')

@app.route('/persiapan-pembuktian')
def persiapan_pembuktian():
    """Persiapan Pembuktian - Automatisasi folder dan BA per perusahaan"""
    return render_template('persiapan_pembuktian.html')

@app.route('/process_comprehensive', methods=['POST'])
def process_comprehensive():
    """Process documents with comprehensive keyword replacement"""
    try:
        # Auto-clean old files (lebih dari 1 jam)
        clean_old_processed_files()
        
        # Get form data
        form_data = request.form.to_dict()
        
        # Check if using master folder or uploaded files
        master_folder_json = form_data.get('master_folder_data')
        uploaded_files = request.files.getlist('template_files')
        
        # Get keywords and deleted documents
        keywords_json = form_data.get('keywords', '{}')
        deleted_docs_json = form_data.get('deleted_documents', '[]')
        keywords_to_delete_rows_json = form_data.get('keywords_to_delete_rows', '[]')
        selected_docs_json = form_data.get('selected_documents', '[]')
        
        try:
            keywords = json.loads(keywords_json)
            deleted_documents = json.loads(deleted_docs_json)
            keywords_to_delete_rows = json.loads(keywords_to_delete_rows_json)
            selected_documents = json.loads(selected_docs_json)
        except:
            keywords = generate_comprehensive_keywords(form_data)
            deleted_documents = []
            keywords_to_delete_rows = []
            selected_documents = []
        
        processed_files = []
        failed_files = []
        
        # Process master folder if provided
        if master_folder_json:
            try:
                master_folder_data = json.loads(master_folder_json)
                available_docs = [doc for doc in master_folder_data['documents'] if doc['available']]
                
                # Debug logging
                print(f"DEBUG: Total available docs: {len(available_docs)}")
                print(f"DEBUG: Available doc IDs: {[doc['id'] for doc in available_docs]}")
                print(f"DEBUG: Selected documents (from frontend): {selected_documents}")
                
                # Filter by selected documents if provided
                if selected_documents:
                    # Match selected codes with document IDs
                    # BA Konsultan: Frontend sends "00", "06", "22_lhp", "27_2" → Backend IDs: "format_00", "format_06", etc.
                    # BA TIMLAK: Frontend sends "DH", "01", "02", "03" → Backend IDs: "timlak_DH", "timlak_01", etc.
                    def matches_selection(doc_id, selected_codes):
                        # Extract code from format_XX or timlak_XX pattern
                        if doc_id.startswith('format_'):
                            code = doc_id[7:]  # Remove 'format_' prefix (BA Konsultan)
                            return code in selected_codes
                        elif doc_id.startswith('timlak_'):
                            code = doc_id[7:]  # Remove 'timlak_' prefix (BA TIMLAK)
                            return code in selected_codes
                        return False
                    
                    available_docs = [doc for doc in available_docs 
                                    if matches_selection(doc['id'], selected_documents)]
                    
                    print(f"DEBUG: After filtering by selection: {len(available_docs)}")
                    print(f"DEBUG: Docs to process: {[doc['id'] for doc in available_docs]}")
                
                if not available_docs:
                    return jsonify({
                        'success': False, 
                        'message': 'Tidak ada dokumen yang dipilih atau tersedia untuk diproses. Silakan pilih dokumen dengan mencentang checkbox di tabel.'
                    })
                
                for doc in available_docs:
                    try:
                        source_path = doc['path']
                        output_filename = doc['name']
                        output_path = os.path.join(PROCESSED_FILES_DIR, output_filename)
                        
                        # Process the document with comprehensive keywords
                        success, result = process_docx_comprehensive(source_path, keywords, output_path, deleted_documents, keywords_to_delete_rows)
                        
                        if success:
                            processed_files.append({
                                'filename': output_filename,
                                'original_filename': doc['name'],
                                'document_type': doc['type'],
                                'replacements': result['total_replacements'],
                                'log_entries': result['log_entries'],
                                'keyword_details': result.get('keyword_details', {})
                            })
                        else:
                            failed_files.append({
                                'filename': doc['name'],
                                'error': result.get('error', 'Unknown error')
                            })
                            
                    except Exception as e:
                        failed_files.append({
                            'filename': doc['name'],
                            'error': str(e)
                        })
                        
            except Exception as e:
                return jsonify({'success': False, 'message': f'Error processing master folder: {str(e)}'})
        
        # Process uploaded files (fallback for backward compatibility)
        elif uploaded_files and any(f.filename != '' for f in uploaded_files):
            for file in uploaded_files:
                if file.filename == '':
                    continue
                    
                if not file.filename.lower().endswith('.docx'):
                    failed_files.append({
                        'filename': file.filename,
                        'error': 'File bukan format .docx'
                    })
                    continue
                
                try:
                    # Save uploaded file temporarily
                    temp_input = os.path.join(PROCESSED_FILES_DIR, f"temp_{file.filename}")
                    file.save(temp_input)
                    
                    # Generate output filename
                    output_filename = file.filename
                    output_path = os.path.join(PROCESSED_FILES_DIR, output_filename)
                    
                    # Process the document with comprehensive keywords
                    success, result = process_docx_comprehensive(temp_input, keywords, output_path, deleted_documents, keywords_to_delete_rows)
                    
                    # Remove temp file
                    os.remove(temp_input)
                    
                    if success:
                        processed_files.append({
                            'filename': output_filename,
                            'original_filename': file.filename,
                            'replacements': result['total_replacements'],
                            'log_entries': result['log_entries']
                        })
                    else:
                        failed_files.append({
                            'filename': file.filename,
                            'error': result.get('error', 'Unknown error')
                        })
                        
                except Exception as e:
                    failed_files.append({
                        'filename': file.filename,
                        'error': str(e)
                    })
        
        else:
            return jsonify({'success': False, 'message': 'Pilih master folder atau upload file template'})
        
        return jsonify({
            'success': True,
            'files': processed_files,
            'failed_files': failed_files,
            'keywords_used': keywords
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/process_keywords', methods=['POST'])
def process_keywords():
    """Process documents with keyword replacement"""
    try:
        # Clear previous results
        shutil.rmtree(PROCESSED_FILES_DIR, ignore_errors=True)
        os.makedirs(PROCESSED_FILES_DIR, exist_ok=True)
        
        # Get form data
        form_data = request.form.to_dict()
        
        # Get uploaded files
        uploaded_files = request.files.getlist('template_files')
        
        if not uploaded_files or all(f.filename == '' for f in uploaded_files):
            return jsonify({'success': False, 'message': 'Tidak ada file yang diupload'})
        
        # Generate keywords from form data
        keywords = generate_keywords_from_form(form_data)
        
        # Process each uploaded file
        processed_files = []
        failed_files = []
        
        for file in uploaded_files:
            if file.filename == '':
                continue
                
            if not file.filename.lower().endswith('.docx'):
                failed_files.append({
                    'filename': file.filename,
                    'error': 'File bukan format .docx'
                })
                continue
            
            try:
                # Save uploaded file temporarily
                temp_input = os.path.join(PROCESSED_FILES_DIR, f"temp_{file.filename}")
                file.save(temp_input)
                
                # Generate output filename
                output_filename = file.filename
                if 'Format' in output_filename:
                    # Replace 'Format' with document type info if available
                    doc_type = form_data.get('document_type', '')
                    if doc_type and doc_type in DOCUMENT_TYPES:
                        output_filename = output_filename.replace('Format', f'{doc_type}-{form_data.get("kode_pokja", "")}-')
                
                output_path = os.path.join(PROCESSED_FILES_DIR, output_filename)
                
                # Process the document
                success, result = process_docx_keywords(temp_input, keywords, output_path)
                
                # Remove temp file
                os.remove(temp_input)
                
                if success:
                    processed_files.append({
                        'filename': output_filename,
                        'original_filename': file.filename,
                        'replacements': result['total_replacements'],
                        'log_entries': result['log_entries']
                    })
                else:
                    failed_files.append({
                        'filename': file.filename,
                        'error': result.get('error', 'Unknown error')
                    })
                    
            except Exception as e:
                failed_files.append({
                    'filename': file.filename,
                    'error': str(e)
                })
        
        return jsonify({
            'success': True,
            'files': processed_files,
            'failed_files': failed_files,
            'keywords_used': keywords
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/download_results')
def download_results():
    """Download all processed files as a ZIP"""
    try:
        # Create ZIP file
        zip_filename = f"BA_Generated_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        zip_path = os.path.join(PROCESSED_FILES_DIR, zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(PROCESSED_FILES_DIR):
                for file in files:
                    if file.endswith('.docx'):  # Only include processed DOCX files
                        file_path = os.path.join(root, file)
                        arcname = file  # Just the filename, not the full path
                        zipf.write(file_path, arcname)
        
        return send_file(zip_path, as_attachment=True, download_name=zip_filename)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download_file/<filename>')
def download_file(filename):
    """Download a specific processed file"""
    try:
        return send_from_directory(PROCESSED_FILES_DIR, filename, as_attachment=True)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/document_types')
def get_document_types():
    """Get available document types"""
    return jsonify({
        'document_types': DOCUMENT_TYPES
    })

@app.route('/api/save_defaults', methods=['POST'])
def save_defaults():
    """Save form defaults to server storage"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'message': 'No data provided'})
        
        # Create saved_data directory if it doesn't exist
        saved_data_dir = os.path.join(os.getcwd(), 'saved_data')
        os.makedirs(saved_data_dir, exist_ok=True)
        
        # Save data with timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"defaults_{timestamp}.json"
        filepath = os.path.join(saved_data_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        return jsonify({
            'success': True,
            'message': 'Data berhasil disimpan',
            'filename': filename
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/load_defaults/<filename>')
def load_defaults(filename):
    """Load form defaults from server storage"""
    try:
        saved_data_dir = os.path.join(os.getcwd(), 'saved_data')
        filepath = os.path.join(saved_data_dir, filename)
        
        if not os.path.exists(filepath):
            return jsonify({'success': False, 'message': 'File tidak ditemukan'})
        
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        return jsonify({
            'success': True,
            'data': data,
            'message': 'Data berhasil dimuat'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/list_saved_defaults')
def list_saved_defaults():
    """List all saved defaults files"""
    try:
        saved_data_dir = os.path.join(os.getcwd(), 'saved_data')
        
        if not os.path.exists(saved_data_dir):
            return jsonify({'success': True, 'files': []})
        
        files = []
        for filename in os.listdir(saved_data_dir):
            if filename.endswith('.json') and filename.startswith('defaults_'):
                filepath = os.path.join(saved_data_dir, filename)
                stat = os.stat(filepath)
                files.append({
                    'filename': filename,
                    'created': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
                    'size': stat.st_size
                })
        
        # Sort by creation time (newest first)
        files.sort(key=lambda x: x['created'], reverse=True)
        
        return jsonify({'success': True, 'files': files})
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/export_defaults/<filename>')
def export_defaults(filename):
    """Export defaults file for download"""
    try:
        saved_data_dir = os.path.join(os.getcwd(), 'saved_data')
        filepath = os.path.join(saved_data_dir, filename)
        
        if not os.path.exists(filepath):
            return jsonify({'error': 'File tidak ditemukan'}), 404
        
        return send_file(filepath, as_attachment=True, download_name=filename)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/import_defaults', methods=['POST'])
def import_defaults():
    """Import defaults from uploaded JSON file"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'No file uploaded'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'No file selected'})
        
        if not file.filename.lower().endswith('.json'):
            return jsonify({'success': False, 'message': 'File harus berformat JSON'})
        
        # Read and validate JSON
        try:
            data = json.load(file)
        except json.JSONDecodeError:
            return jsonify({'success': False, 'message': 'File JSON tidak valid'})
        
        # Save imported data
        saved_data_dir = os.path.join(os.getcwd(), 'saved_data')
        os.makedirs(saved_data_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"imported_{timestamp}.json"
        filepath = os.path.join(saved_data_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        return jsonify({
            'success': True,
            'data': data,
            'message': 'Data berhasil diimport',
            'filename': filename
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/list_folder_files', methods=['POST'])
def list_folder_files():
    """Quick listing of Word files in POKJA KONSULTAN folder"""
    try:
        data = request.get_json()
        folder_path = data.get('folder_path', '')
        
        if not folder_path:
            return jsonify({'success': False, 'message': 'Path folder tidak valid'})
        
        if not os.path.exists(folder_path):
            return jsonify({'success': False, 'message': 'Folder tidak ditemukan'})
        
        if not os.path.isdir(folder_path):
            return jsonify({'success': False, 'message': 'Path bukan folder'})
        
        # List all .docx files
        available_files = []
        try:
            for file in os.listdir(folder_path):
                if file.lower().endswith('.docx') and not file.startswith('~'):
                    # Extract document code from filename (e.g., "00", "06", "10")
                    # Expected format: "XX. Name.docx" or "XX-Name.docx"
                    import re
                    
                    # Try to match patterns like "00.", "06.", "10.", "22-LHP", "27-1", "27-2"
                    match = re.match(r'^(\d{2})(?:[-\.]|$)', file)
                    if match:
                        code = match.group(1)
                        
                        # Check if there's additional suffix like "22-LHP", "27-1", "27-2"
                        suffix_match = re.match(r'^\d{2}[-\.]([A-Za-z0-9]+)', file)
                        if suffix_match:
                            suffix = suffix_match.group(1).lower()
                            code = f"{code}_{suffix}"  # e.g., "22_lhp", "27_1", "27_2"
                        
                        format_id = f"format_{code}"
                        
                        available_files.append({
                            'code': code,
                            'format_id': format_id,
                            'filename': file,
                            'exists': True
                        })
        except Exception as e:
            return jsonify({'success': False, 'message': f'Error membaca folder: {str(e)}'})
        
        return jsonify({
            'success': True,
            'folder_path': folder_path,
            'files': available_files,
            'total_files': len(available_files)
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/list_folder_files_timlak', methods=['POST'])
def list_folder_files_timlak():
    """Quick listing of Word files in TIMLAK folder - handles special DH document"""
    try:
        data = request.get_json()
        folder_path = data.get('folder_path', '')
        
        if not folder_path:
            return jsonify({'success': False, 'message': 'Path folder tidak valid'})
        
        if not os.path.exists(folder_path):
            return jsonify({'success': False, 'message': 'Folder tidak ditemukan'})
        
        if not os.path.isdir(folder_path):
            return jsonify({'success': False, 'message': 'Path bukan folder'})
        
        # List all .docx files
        available_files = []
        try:
            for file in os.listdir(folder_path):
                if file.lower().endswith('.docx') and not file.startswith('~'):
                    # Special handling for TIMLAK documents
                    code = None
                    format_id = None
                    
                    # Check if it's the special "!Daftar Hadir.docx" file (DH)
                    if file.startswith('!Daftar Hadir'):
                        code = 'DH'
                        format_id = 'format_DH'
                    else:
                        # Extract document code from filename (e.g., "01", "02", "03")
                        # Expected format: "XX. Name.docx" or "XX-Name.docx"
                        import re
                        match = re.match(r'^(\d{2})', file)
                        if match:
                            code = match.group(1)
                            format_id = f"format_{code}"
                    
                    if code and format_id:
                        available_files.append({
                            'code': code,
                            'format_id': format_id,
                            'filename': file,
                            'exists': True
                        })
        except Exception as e:
            return jsonify({'success': False, 'message': f'Error membaca folder: {str(e)}'})
        
        return jsonify({
            'success': True,
            'folder_path': folder_path,
            'files': available_files,
            'total_files': len(available_files)
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/validate_master_pokja_konsultan', methods=['POST'])
def validate_master_folder():
    """Validate master folder and return available documents"""
    try:
        data = request.get_json()
        folder_path = data.get('folder_path', '')
        
        if not folder_path or not os.path.exists(folder_path):
            return jsonify({'success': False, 'message': 'Folder tidak ditemukan'})
        
        # Expected document list based on image attachment
        expected_documents = [
            {'id': 'format_00', 'name': '00-Format-Cover.docx', 'type': 'Cover'},
            {'id': 'format_06', 'name': '06-Format-BA Pemberian Penjelasan Kualifikasi.docx', 'type': 'BA Pemberian Penjelasan Kualifikasi'},
            {'id': 'format_10', 'name': '10-Format-BA Hasil Evaluasi Kualifikasi.docx', 'type': 'BA Hasil Evaluasi Kualifikasi'},
            {'id': 'format_11', 'name': '11-Format-BA Penetapan Daftar Pendek.docx', 'type': 'BA Penetapan Daftar Pendek'},
            {'id': 'format_12', 'name': '12-Format-Pengumuman Daftar Pendek.docx', 'type': 'Pengumuman Daftar Pendek'},
            {'id': 'format_13', 'name': '13-Format-BA Jawab Sanggah PQ.docx', 'type': 'BA Jawab Sanggah PQ'},
            {'id': 'format_14', 'name': '14-Format-BA Pemberian Penjelasan Seleksi.docx', 'type': 'BA Pemberian Penjelasan Seleksi'},
            {'id': 'format_17', 'name': '17-Format-BA Admin Teknis File I.docx', 'type': 'BA Admin Teknis File I'},
            {'id': 'format_19', 'name': '19-Format-BA Evaluasi Biaya.docx', 'type': 'BA Evaluasi Biaya'},
            {'id': 'format_20', 'name': '20-Format-BA Kombinasi Teknis dan Biaya.docx', 'type': 'BA Kombinasi Teknis dan Biaya'},
            {'id': 'format_21', 'name': '21-Format-Surat Klarifikasi Personel.docx', 'type': 'Surat Klarifikasi Personel'},
            {'id': 'format_22', 'name': '22-Format-BA KLARIFIKASI PENETAPAN PEMENANG.docx', 'type': 'BA Klarifikasi Penetapan Pemenang'},
            {'id': 'format_22_lhp', 'name': '22-Format----------LHP.docx', 'type': 'LHP'},
            {'id': 'format_24', 'name': '24-Format-BA Penetapan Pemenang.docx', 'type': 'BA Penetapan Pemenang'},
            {'id': 'format_25', 'name': '25-Format-BA Pengumuman Pemenang.docx', 'type': 'BA Pengumuman Pemenang'},
            {'id': 'format_26', 'name': '26-Format-BA Jawab Sanggah Seleksi.docx', 'type': 'BA Jawab Sanggah Seleksi'},
            {'id': 'format_27_1', 'name': '27-1-Format-BA Klarifikasi Negosiasi.docx', 'type': 'BA Klarifikasi Negosiasi'},
            {'id': 'format_27_2', 'name': '27-2-Format-Daftar Hadir Klarneg.docx', 'type': 'Daftar Hadir Klarneg'},
            {'id': 'format_28', 'name': '28-Format-BAHP.docx', 'type': 'BAHP'},
            {'id': 'format_29', 'name': '29-Format-Surat Penyampaian BAHP.docx', 'type': 'Surat Penyampaian BAHP'},
            {'id': 'format_96', 'name': '96-Format-Surat pernyataan Klarifikasi personil dan paket 1 dan 2.docx', 'type': 'Surat Pernyataan Klarifikasi'},
            {'id': 'format_97', 'name': '97-Format-BA Seleksi Ulang.docx', 'type': 'BA Seleksi Ulang'},
            {'id': 'format_99', 'name': '99-Format-TTD Pokja.docx', 'type': 'TTD Pokja'}
        ]
        
        # Scan folder for .docx files
        available_files = []
        if os.path.isdir(folder_path):
            for file in os.listdir(folder_path):
                if file.lower().endswith('.docx') and not file.startswith('~'):
                    available_files.append(file)
        
        # Check which documents are available
        validated_documents = []
        for doc in expected_documents:
            is_available = doc['name'] in available_files
            validated_documents.append({
                'id': doc['id'],
                'name': doc['name'],
                'type': doc['type'],
                'available': is_available,
                'path': os.path.join(folder_path, doc['name']) if is_available else None
            })
        
        return jsonify({
            'success': True,
            'folder_path': folder_path,
            'documents': validated_documents,
            'total_expected': len(expected_documents),
            'total_available': sum(1 for doc in validated_documents if doc['available'])
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/get_folder_documents', methods=['POST'])
def get_folder_documents():
    """Get all .docx files from a folder"""
    try:
        data = request.get_json()
        folder_path = data.get('folder_path', '')
        
        if not folder_path or not os.path.exists(folder_path):
            return jsonify({'success': False, 'message': 'Folder tidak ditemukan'})
        
        documents = []
        if os.path.isdir(folder_path):
            for file in os.listdir(folder_path):
                if file.lower().endswith('.docx') and not file.startswith('~'):
                    file_path = os.path.join(folder_path, file)
                    documents.append({
                        'name': file,
                        'path': file_path,
                        'size': os.path.getsize(file_path)
                    })
        
        return jsonify({
            'success': True,
            'documents': documents,
            'count': len(documents)
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

@app.route('/api/validate_master_pembuktian', methods=['POST'])
def validate_master_pembuktian():
    """Validate master folder for pembuktian contains required files"""
    try:
        data = request.json
        folder_path = data.get('folder_path', '')
        
        if not folder_path:
            return jsonify({
                'success': False,
                'error': 'Folder path tidak boleh kosong'
            })
        
        if not os.path.exists(folder_path):
            return jsonify({
                'success': False,
                'error': 'Folder tidak ditemukan'
            })
        
        # Required files
        required_files = [
            {
                'name': '09.no-1-BA Pembuktian.docx',
                'type': 'BA',
                'description': 'Template Berita Acara Pembuktian'
            },
            {
                'name': '09.no-3-Lamp Kerja Sejenis.xlsx',
                'type': 'Excel',
                'description': 'Template Lampiran Pengalaman Sejenis'
            },
            {
                'name': '09.no-4-Daftar Hadir Pembuktian.docx',
                'type': 'Daftar Hadir',
                'description': 'Template Daftar Hadir Pembuktian'
            }
        ]
        
        # Check each file
        validation_results = []
        for file_info in required_files:
            file_path = os.path.join(folder_path, file_info['name'])
            found = os.path.exists(file_path)
            
            validation_results.append({
                'name': file_info['name'],
                'type': file_info['type'],
                'description': file_info['description'],
                'found': found,
                'path': file_path if found else None
            })
        
        # Check if all files found
        all_found = all(f['found'] for f in validation_results)
        
        return jsonify({
            'success': True,
            'all_valid': all_found,
            'files': validation_results,
            'folder_path': folder_path
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })

@app.route('/api/validate_master_timlak_konsultan', methods=['POST'])
def validate_master_timlak():
    """Validate Master BA Timlak Konsultan folder and return available documents"""
    try:
        data = request.get_json()
        folder_path = data.get('folder_path', '')
        
        if not folder_path or not os.path.exists(folder_path):
            return jsonify({'success': False, 'message': 'Folder tidak ditemukan'})
        
        # Expected TIMLAK documents (9 documents: DH + 01-08)
        expected_timlak_documents = [
            {'id': 'timlak_DH', 'name': '!Daftar Hadir.docx', 'type': 'Daftar Hadir', 'doc_num': 'DH'},
            {'id': 'timlak_01', 'name': '01. BA Reviu Persiapan Pengadaan.docx', 'type': 'BA Reviu Persiapan Pengadaan', 'doc_num': '01'},
            {'id': 'timlak_02', 'name': '02. Memo Dinas BA Reviu Persiapan Pengadaan.docx', 'type': 'Memo Dinas BA Reviu Persiapan Pengadaan', 'doc_num': '02'},
            {'id': 'timlak_03', 'name': '03. Surat Penetapan BA Persiapan Pengadaan PPK.docx', 'type': 'Surat Penetapan BA Persiapan Pengadaan PPK', 'doc_num': '03'},
            {'id': 'timlak_04', 'name': '04. BA Reviu Dokumen Kualifikasi.docx', 'type': 'BA Reviu Dokumen Kualifikasi', 'doc_num': '04'},
            {'id': 'timlak_05', 'name': '05. BA Reviu Dokumen Seleksi.docx', 'type': 'BA Reviu Dokumen Seleksi', 'doc_num': '05'},
            {'id': 'timlak_06', 'name': '06. Catatan Pemeriksaan BA Hasil Reviu Dokumen Pemilihan.docx', 'type': 'Catatan Pemeriksaan BA Hasil Reviu Dokumen Pemilihan', 'doc_num': '06'},
            {'id': 'timlak_07', 'name': '07. Nota Dinas Penetapan Dokumen Pemilihan.docx', 'type': 'Nota Dinas Penetapan Dokumen Pemilihan', 'doc_num': '07'},
            {'id': 'timlak_08', 'name': '08. BA Penetapan Dokumen Pemilihan.docx', 'type': 'BA Penetapan Dokumen Pemilihan', 'doc_num': '08'}
        ]
        
        # Scan folder for .docx files
        available_files = {}
        if os.path.isdir(folder_path):
            for file in os.listdir(folder_path):
                if file.lower().endswith('.docx') and not file.startswith('~'):
                    available_files[file] = True
        
        # Check which documents are available
        validated_documents = []
        for doc in expected_timlak_documents:
            is_available = doc['name'] in available_files
            validated_documents.append({
                'id': doc['id'],
                'name': doc['name'],
                'type': doc['type'],
                'doc_num': doc['doc_num'],
                'available': is_available,
                'path': os.path.join(folder_path, doc['name']) if is_available else None
            })
        
        return jsonify({
            'success': True,
            'folder_path': folder_path,
            'documents': validated_documents,
            'total_expected': len(expected_timlak_documents),
            'total_available': sum(1 for doc in validated_documents if doc['available'])
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

def fill_excel_pengalaman(excel_path, company_data, pengalaman_data, form_data):
    """
    Fill Excel template - PRESERVE template structure, only:
    1. Delete unused sheets (Sheet1/Sheet2/Sheet3 based on KSO)
    2. Hide excess rows (template has max 15 sejenis, 10 beda jenis)
    3. Replace placeholders globally
    
    Template Structure:
    - Row 7-21: Sejenis data (15 rows max)
    - Row 23-32: Beda jenis data (10 rows max)
    
    Args:
        excel_path: Path to Excel file (09.no-3-Lamp Kerja Sejenis.xlsx)
        company_data: {"no": 1, "leadfirm": "PT. XXX", "kso_anggota2": "PT. KSO1", "kso_anggota3": "PT. KSO2"}
        pengalaman_data: {"sejenis": 7, "tahun_sejenis": 10, "beda_jenis": 4, "tahun_beda_jenis": 4}
        form_data: {all form fields including kode_pokja, nama_paket, etc.}
    """
    try:
        wb = load_workbook(excel_path)
        
        # ============================================================
        # STEP 1: DETECT KSO & DELETE UNUSED SHEETS
        # ============================================================
        kso_anggota2 = company_data.get('kso_anggota2', '').strip()
        kso_anggota3 = company_data.get('kso_anggota3', '').strip()
        
        print(f"[EXCEL] KSO Detection:")
        print(f"  - kso_anggota2: '{kso_anggota2}'")
        print(f"  - kso_anggota3: '{kso_anggota3}'")
        
        if not kso_anggota2:
            selected_sheet_name = 'Sheet1'
            sheets_to_delete = ['Sheet2', 'Sheet3']
            print(f"  - Decision: NO KSO → Sheet1 (lead firm only)")
        elif not kso_anggota3:
            selected_sheet_name = 'Sheet2'
            sheets_to_delete = ['Sheet1', 'Sheet3']
            print(f"  - Decision: 2 anggota → Sheet2")
        else:
            selected_sheet_name = 'Sheet3'
            sheets_to_delete = ['Sheet1', 'Sheet2']
            print(f"  - Decision: 3 anggota → Sheet3")
        
        print(f"  - Selected sheet: {selected_sheet_name}")
        print(f"  - Sheets to delete: {sheets_to_delete}")
        
        for sheet_name in sheets_to_delete:
            if sheet_name in wb.sheetnames:
                print(f"  - Deleting sheet: {sheet_name}")
                del wb[sheet_name]
        
        print(f"  - Remaining sheets: {wb.sheetnames}")
        
        ws = wb[selected_sheet_name]
        
        # ============================================================
        # STEP 2: GET PENGALAMAN COUNTS
        # ============================================================
        sejenis_count = int(pengalaman_data.get('sejenis', 7))
        tahun_sejenis = int(pengalaman_data.get('tahun_sejenis', 10))
        beda_jenis_count = int(pengalaman_data.get('beda_jenis', 4))
        tahun_beda_jenis = int(pengalaman_data.get('tahun_beda_jenis', 4))
        
        # Template constants (fixed positions in template)
        SEJENIS_START_ROW = 7
        SEJENIS_MAX_ROWS = 15
        BEDA_JENIS_START_ROW = 23  # Original position before deletion
        BEDA_JENIS_MAX_ROWS = 10
        
        # ============================================================
        # STEP 3: HIDE EXCESS SEJENIS ROWS
        # ============================================================
        # Always unhide full sejenis block first, then hide rows after count
        for r in range(SEJENIS_START_ROW, SEJENIS_START_ROW + SEJENIS_MAX_ROWS):
            ws.row_dimensions[r].hidden = False
        if sejenis_count < 0:
            sejenis_count = 0
        if sejenis_count > SEJENIS_MAX_ROWS:
            sejenis_count = SEJENIS_MAX_ROWS
        # Hide rows exceeding the requested count
        for r in range(SEJENIS_START_ROW + sejenis_count, SEJENIS_START_ROW + SEJENIS_MAX_ROWS):
            ws.row_dimensions[r].hidden = True
        
        # ============================================================
        # STEP 4: HIDE EXCESS BEDA JENIS ROWS
        # ============================================================
        # Beda jenis block is fixed at 23-32 regardless of hidden sejenis
        for r in range(BEDA_JENIS_START_ROW, BEDA_JENIS_START_ROW + BEDA_JENIS_MAX_ROWS):
            ws.row_dimensions[r].hidden = False
        if beda_jenis_count < 0:
            beda_jenis_count = 0
        if beda_jenis_count > BEDA_JENIS_MAX_ROWS:
            beda_jenis_count = BEDA_JENIS_MAX_ROWS
        for r in range(BEDA_JENIS_START_ROW + beda_jenis_count, BEDA_JENIS_START_ROW + BEDA_JENIS_MAX_ROWS):
            ws.row_dimensions[r].hidden = True
        
        # ============================================================
        # STEP 5: REPLACE ALL PLACEHOLDERS
        # ============================================================
        # Get all data for replacement
        leadfirm = company_data.get('leadfirm', '')
        nama_kso = company_data.get('nama_kso', '')
        # note_pengalaman may come either at top-level or under 'keywords'
        note_pengalaman = (
            (form_data.get('note_pengalaman') if isinstance(form_data, dict) else None)
            or (form_data.get('keywords', {}).get('note_pengalaman') if isinstance(form_data, dict) else None)
            or ''
        )
        
        print(f"[EXCEL] Replacing placeholders:")
        print(f"  - leadfirm: {leadfirm}")
        print(f"  - nama_kso: {nama_kso}")
        print(f"  - note_pengalaman: {note_pengalaman}")
        print(f"  - kso_anggota2: {kso_anggota2}")
        print(f"  - kso_anggota3: {kso_anggota3}")
        
        # Replace in ALL cells (global replacement)
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    # Replace all placeholders
                    cell.value = (cell.value
                        .replace('{X_tahun_sejenis}', str(tahun_sejenis))
                        .replace('{X_tahun_beda_jenis}', str(tahun_beda_jenis))
                        .replace('{note_pengalaman}', note_pengalaman)
                        .replace('{leadfirm}', leadfirm)
                        .replace('{nama_kso}', nama_kso)
                        .replace('{kso_anggota2}', kso_anggota2)
                        .replace('{kso_anggota3}', kso_anggota3)
                    )
        
        # ============================================================
        # STEP 6: SAVE WORKBOOK
        # ============================================================
        wb.save(excel_path)
        wb.close()
        
        return True
        
    except Exception as e:
        print(f"Error filling Excel: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

@app.route('/api/generate_pembuktian_folders', methods=['POST'])
def generate_pembuktian_folders():
    """
    Generate folders and documents for multiple companies (batch processing)
    
    Expected JSON:
    {
      "companies": [
        {"no": 1, "name": "PT. XXX", "kso": ["PT. KSO1", "PT. KSO2"]},
        {"no": 2, "name": "PT. YYY", "kso": []}
      ],
      "pengalaman": {
        "sejenis": 7,
        "tahun_sejenis": 10,
        "beda_jenis": 4,
        "tahun_beda_jenis": 4
      },
      "master_folder": "C:\\...\\Master Pembuktian"
    }
    
    Returns JSON:
    {
      "success": true,
      "message": "13 folder berhasil dibuat!",
      "download_url": "/download_file/Pembuktian_20240117_153045.zip"
    }
    """
    try:
        data = request.json
        companies = data.get('companies', [])
        pengalaman = data.get('pengalaman', {})
        master_folder = data.get('master_folder', '')
        
        # Validation
        if not companies:
            return jsonify({
                'success': False,
                'error': 'Tidak ada perusahaan yang diproses'
            }), 400
        
        if not master_folder or not os.path.exists(master_folder):
            return jsonify({
                'success': False,
                'error': 'Folder master data tidak valid'
            }), 400
        
        # ===== CLEANUP OLD FILES =====
        # Only remove previous Pembuktian_* outputs to avoid interfering with other features
        os.makedirs(PROCESSED_FILES_DIR, exist_ok=True)
        
        print(f"[CLEANUP] Scanning {PROCESSED_FILES_DIR} for old Pembuktian_* files...")
        cleanup_count = 0
        try:
            for name in os.listdir(PROCESSED_FILES_DIR):
                if name.startswith('Pembuktian_'):
                    path = os.path.join(PROCESSED_FILES_DIR, name)
                    try:
                        if os.path.isdir(path):
                            print(f"[CLEANUP] Removing folder: {name}")
                            shutil.rmtree(path, ignore_errors=True)
                            cleanup_count += 1
                        elif os.path.isfile(path):
                            print(f"[CLEANUP] Removing file: {name}")
                            os.remove(path)
                            cleanup_count += 1
                    except Exception as e:
                        # Best-effort cleanup; continue even if a file is locked
                        print(f"[CLEANUP] Warning: Could not remove {name}: {str(e)}")
                        pass
        except FileNotFoundError:
            print(f"[CLEANUP] Directory not found, creating: {PROCESSED_FILES_DIR}")
            os.makedirs(PROCESSED_FILES_DIR, exist_ok=True)
        
        print(f"[CLEANUP] Cleaned up {cleanup_count} old Pembuktian_* items")
        
        # Create timestamp for unique folder name
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_folder = os.path.join(PROCESSED_FILES_DIR, f'Pembuktian_{timestamp}')
        os.makedirs(output_folder, exist_ok=True)
        
        # Required master files
        master_files = [
            '09.no-1-BA Pembuktian.docx',
            '09.no-3-Lamp Kerja Sejenis.xlsx',
            '09.no-4-Daftar Hadir Pembuktian.docx'
        ]
        
        # Process each company
        for company in companies:
            company_no = company.get('no', 0)
            company_name = company.get('name', '')
            kso_list = company.get('kso', [])
            nama_kso = company.get('namaKSO', '')  # Nama resmi KSO (berbeda dari anggota)
            
            # Format company number with leading zero: 01, 02, 03, etc.
            formatted_no = str(company_no).zfill(2)
            
            # Create company folder: "01- PT. Company Name", "02- CV. Company", etc.
            folder_name = f"{formatted_no}- {company_name}"
            company_folder = os.path.join(output_folder, folder_name)
            os.makedirs(company_folder, exist_ok=True)
            
            # Copy 3 master files to company folder with new naming format
            for file_name in master_files:
                src_path = os.path.join(master_folder, file_name)
                
                # Generate new file name with company number: 09.01-1-BA Pembuktian.docx
                # Replace "09.no-" with "09.{formatted_no}-"
                new_file_name = file_name.replace('09.no-', f'09.{formatted_no}-')
                dst_path = os.path.join(company_folder, new_file_name)
                
                if not os.path.exists(src_path):
                    return jsonify({
                        'success': False,
                        'error': f'File master tidak ditemukan: {file_name}'
                    }), 400
                
                # Copy file
                shutil.copy2(src_path, dst_path)
                
                # Edit Excel file to generate rows based on pengalaman counts
                if file_name == '09.no-3-Lamp Kerja Sejenis.xlsx':
                    # Prepare company data with KSO info
                    company_info = {
                        'no': company_no,
                        'name': company_name,
                        'nama_kso': nama_kso,  # Nama resmi KSO
                        'leadfirm': company_name,
                        'kso_anggota2': kso_list[0] if len(kso_list) > 0 else '',
                        'kso_anggota3': kso_list[1] if len(kso_list) > 1 else '',
                        'kso_anggota4': kso_list[2] if len(kso_list) > 2 else '',
                        'kso_anggota5': kso_list[3] if len(kso_list) > 3 else ''
                    }
                    
                    # Fill Excel with generated rows
                    fill_excel_pengalaman(dst_path, company_info, pengalaman, data)
                
                # Edit Word documents (BA Pembuktian and Daftar Hadir)
                if file_name.endswith('.docx'):
                    # Prepare keywords for Word replacement
                    formatted_no = str(company_no).zfill(2)  # 01, 02, 03, dst.
                    word_keywords = {
                        'no': formatted_no,  # Nomor urut perusahaan
                        'nama_kso': nama_kso,
                        'leadfirm': company_name,
                        'kso_anggota2': kso_list[0] if len(kso_list) > 0 else '',
                        'kso_anggota3': kso_list[1] if len(kso_list) > 1 else '',
                        'kso_anggota4': kso_list[2] if len(kso_list) > 2 else '',
                        'kso_anggota5': kso_list[3] if len(kso_list) > 3 else ''
                    }
                    
                    # Add global keywords (from form data)
                    global_keywords = data.get('keywords', {})
                    word_keywords.update(global_keywords)
                    
                    # Identify empty KSO placeholders for row deletion
                    # If a KSO member is empty, mark it for row deletion
                    keywords_to_delete_rows = []
                    if not kso_list or len(kso_list) == 0:
                        keywords_to_delete_rows.extend(['kso_anggota2', 'kso_anggota3', 'kso_anggota4', 'kso_anggota5'])
                    else:
                        if len(kso_list) < 1:
                            keywords_to_delete_rows.append('kso_anggota2')
                        if len(kso_list) < 2:
                            keywords_to_delete_rows.append('kso_anggota3')
                        if len(kso_list) < 3:
                            keywords_to_delete_rows.append('kso_anggota4')
                        if len(kso_list) < 4:
                            keywords_to_delete_rows.append('kso_anggota5')
                    
                    # Process Word document with comprehensive replacement and row deletion
                    process_docx_comprehensive(dst_path, word_keywords, dst_path, 
                                             keywords_to_delete_rows=keywords_to_delete_rows)
        
        # Create ZIP file
        zip_filename = f'Pembuktian_{timestamp}.zip'
        zip_path = os.path.join(PROCESSED_FILES_DIR, zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(output_folder):
                for file in files:
                    file_path = os.path.join(root, file)
                    # Create relative path for ZIP (preserve folder structure)
                    arcname = os.path.relpath(file_path, output_folder)
                    zipf.write(file_path, arcname)
        
        # Clean up temporary output folder
        shutil.rmtree(output_folder)
        
        # Prepare detailed response with keyword tracking
        keywords = data.get('keywords', {})
        
        # Build keyword details for all files
        keyword_details = {}
        for key, value in keywords.items():
            if value:  # Only count non-empty values
                # Estimate replacement count based on file types:
                # - Excel: pengalaman keywords appear once
                # - Word (2 files): all keywords appear multiple times
                if key.startswith('pengalaman') or key.startswith('X_tahun') or key == 'note_pengalaman':
                    count_per_company = 1  # Excel only
                else:
                    count_per_company = 3  # Word files (BA + Daftar Hadir, multiple locations)
                
                keyword_details[key] = {
                    'value': str(value),
                    'count': count_per_company * len(companies)
                }
        
        # Add KSO-related placeholders (per-company basis)
        # These appear in both BA Pembuktian.docx and Excel
        kso_placeholders = ['no', 'nama_kso', 'leadfirm', 'kso_anggota2', 'kso_anggota3', 'kso_anggota4', 'kso_anggota5']
        for placeholder in kso_placeholders:
            keyword_details[placeholder] = {
                'value': f'(berbeda per perusahaan)',
                'count': len(companies) * 5  # Estimate: Excel (1x) + BA.docx (3x) + Daftar Hadir.docx (1x)
            }
        
        return jsonify({
            'success': True,
            'message': f'✓ {len(companies)} folder berhasil dibuat dengan {len(master_files)} file per folder',
            'download_url': f'/download_file/{zip_filename}',
            'companies_processed': len(companies),
            'pengalaman': pengalaman,
            'files': [
                {
                    'filename': '09.XX-1-BA Pembuktian.docx',
                    'replacements': sum(detail['count'] for detail in keyword_details.values() if not detail['value'].startswith('(berbeda')),
                    'keyword_details': keyword_details
                },
                {
                    'filename': '09.XX-3-Lamp Kerja Sejenis.xlsx (Auto-Generated Rows)',
                    'replacements': len(companies) * (pengalaman.get('sejenis', 0) + pengalaman.get('beda_jenis', 0)),
                    'keyword_details': {
                        'pengalaman_sejenis': {
                            'value': f"{pengalaman.get('sejenis', 0)} pengalaman",
                            'count': len(companies)
                        },
                        'pengalaman_beda_jenis': {
                            'value': f"{pengalaman.get('beda_jenis', 0)} pengalaman",
                            'count': len(companies)
                        }
                    }
                },
                {
                    'filename': '09.XX-4-Daftar Hadir Pembuktian.docx',
                    'replacements': sum(1 for p in kso_placeholders for _ in range(len(companies))),
                    'keyword_details': {p: keyword_details[p] for p in kso_placeholders}
                }
            ]
        })
        
    except Exception as e:
        # Log detailed error for debugging
        print(f"[ERROR] Exception in generate_pembuktian_folders:")
        print(f"  Type: {type(e).__name__}")
        print(f"  Message: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # Clean up on error
        try:
            if 'output_folder' in locals() and os.path.exists(output_folder):
                print(f"[CLEANUP] Removing incomplete output folder: {output_folder}")
                shutil.rmtree(output_folder, ignore_errors=True)
            if 'zip_path' in locals() and os.path.exists(zip_path):
                print(f"[CLEANUP] Removing incomplete zip file: {zip_path}")
                os.remove(zip_path)
        except Exception as cleanup_error:
            print(f"[CLEANUP] Error during cleanup: {str(cleanup_error)}")
            pass
        
        return jsonify({
            'success': False,
            'error': f'Terjadi kesalahan: {str(e)}'
        }), 500

@app.route('/api/load_pokja_members')
def load_pokja_members():
    """Load POKJA members from CSV file with optional group filter"""
    try:
        csv_path = os.path.join(os.getcwd(), 'pokja_members.csv')
        
        if not os.path.exists(csv_path):
            return jsonify({
                'success': False,
                'error': 'File pokja_members.csv tidak ditemukan'
            }), 404
        
        # Get group filter from query parameter (konsultan or timlak)
        group_filter = request.args.get('group', 'konsultan')  # Default to konsultan
        
        members = []
        import csv
        
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                # Filter by group if column exists
                if 'group' in row and row['group'].strip() != group_filter:
                    continue
                    
                members.append({
                    'nama': row['nama'].strip(),
                    'nip': row['nip'].strip(),
                    'email': row['email'].strip()
                })
        
        return jsonify({
            'success': True,
            'members': members,
            'group': group_filter
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Error membaca CSV: {str(e)}'
        }), 500

@app.route('/api/preview_document', methods=['POST'])
def preview_document():
    """Generate HTML preview of document with replaced keywords"""
    try:
        import mammoth
        
        data = request.json
        doc_code = data.get('doc_code', '')
        keywords = data.get('keywords', {})
        master_folder = data.get('master_folder', '')
        
        if not doc_code or not master_folder:
            return jsonify({
                'success': False,
                'error': 'Missing doc_code or master_folder'
            }), 400
        
        print(f"[PREVIEW] Doc code: {doc_code}, Master folder: {master_folder}")
        
        # Normalize doc_code (replace underscore with hyphen)
        normalized_code = doc_code.replace('_', '-')
        
        # Special mapping for files that don't follow standard naming
        special_files = {
            'DH': ['!Daftar Hadir.docx', 'Daftar Hadir.docx', '!DH.docx'],
            '22-lhp': ['22-Format----------LHP.docx', '22-LHP.docx', '22-Format-LHP.docx'],
            '22_lhp': ['22-Format----------LHP.docx', '22-LHP.docx', '22-Format-LHP.docx']
        }
        
        # List all files in master folder
        try:
            all_files = os.listdir(master_folder)
            print(f"[PREVIEW] Files in master folder: {all_files[:20]}")  # First 20 files
        except Exception as e:
            print(f"[PREVIEW] Error listing folder: {str(e)}")
            return jsonify({
                'success': False,
                'error': f'Cannot access master folder: {str(e)}'
            }), 500
        
        # Find file that starts with doc_code or matches special mapping
        doc_path = None
        
        # First check special mappings (case insensitive)
        doc_code_lower = doc_code.lower()
        normalized_lower = normalized_code.lower()
        
        if doc_code_lower in special_files or normalized_lower in special_files:
            special_names = special_files.get(doc_code_lower) or special_files.get(normalized_lower)
            for special_name in special_names:
                if special_name in all_files:
                    doc_path = os.path.join(master_folder, special_name)
                    print(f"[PREVIEW] Found special file: {special_name}")
                    break
        
        # If not found in special mapping, search by prefix
        if not doc_path:
            for filename in all_files:
                if not filename.lower().endswith('.docx'):
                    continue
                
                # Check if filename starts with the doc code
                if filename.startswith(doc_code) or filename.startswith(normalized_code):
                    doc_path = os.path.join(master_folder, filename)
                    print(f"[PREVIEW] Found matching file: {filename}")
                    break
        
        if not doc_path:
            error_msg = f'Dokumen tidak ditemukan untuk code: {doc_code}\n'
            error_msg += f'Master folder: {master_folder}\n'
            error_msg += f'File yang tersedia dimulai dengan: {", ".join([f.split("-")[0] for f in all_files if f.endswith(".docx")])}'
            print(f"[PREVIEW] ERROR: {error_msg}")
            return jsonify({
                'success': False,
                'error': error_msg
            }), 404
        
        # Convert DOCX to HTML using mammoth
        with open(doc_path, 'rb') as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value
            messages = result.messages  # Any warnings during conversion
        
        # Replace keywords in HTML
        for key, value in keywords.items():
            placeholder = '{' + key + '}'
            html_content = html_content.replace(placeholder, str(value) if value else '')
        
        # Add styling to make it look better
        styled_html = f"""
        <style>
            .docx-preview {{
                font-family: 'Calibri', 'Arial', sans-serif;
                font-size: 11pt;
                line-height: 1.5;
                color: #000;
                padding: 20px;
                background: white;
            }}
            .docx-preview p {{
                margin: 6pt 0;
            }}
            .docx-preview table {{
                border-collapse: collapse;
                width: 100%;
                margin: 10px 0;
            }}
            .docx-preview table td, .docx-preview table th {{
                border: 1px solid #000;
                padding: 5px;
            }}
            .docx-preview strong {{
                font-weight: bold;
            }}
            .docx-preview em {{
                font-style: italic;
            }}
            .docx-preview u {{
                text-decoration: underline;
            }}
        </style>
        <div class="docx-preview">
            {html_content}
        </div>
        """
        
        # Check for any remaining placeholders
        remaining_placeholders = re.findall(r'\{([a-zA-Z0-9_\-]+)\}', html_content)
        
        return jsonify({
            'success': True,
            'doc_code': doc_code,
            'html': styled_html,
            'warnings': [str(m) for m in messages] if messages else [],
            'remaining_placeholders': list(set(remaining_placeholders)),
            'doc_path': doc_path
        })
        
    except ImportError:
        return jsonify({
            'success': False,
            'error': 'Library mammoth belum terinstall. Jalankan: pip install mammoth'
        }), 500
    except Exception as e:
        print(f"[ERROR] Preview document error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Terjadi kesalahan: {str(e)}'
        }), 500

@app.route('/api/preview_excel', methods=['POST'])
def preview_excel():
    """Generate HTML preview of Excel with replaced keywords AND processed structure (hide rows, select sheet)"""
    try:
        from openpyxl import load_workbook
        from openpyxl.utils import get_column_letter
        import tempfile
        import shutil
        
        data = request.json
        file_path = data.get('file_path', '')
        keywords = data.get('keywords', {})
        company_data = data.get('company_data', {})
        pengalaman_data = data.get('pengalaman_data', {})
        
        print(f"[DEBUG] Preview Excel requested for: {file_path}")
        print(f"[DEBUG] Keywords count: {len(keywords)}")
        print(f"[DEBUG] Company data: {company_data}")
        print(f"[DEBUG] Pengalaman data: {pengalaman_data}")
        
        if not file_path or not os.path.exists(file_path):
            error_msg = f'File Excel tidak ditemukan: {file_path}'
            print(f"[ERROR] {error_msg}")
            return jsonify({
                'success': False,
                'error': error_msg
            }), 404
        
        # Create temporary copy for processing (don't modify original template)
        temp_dir = tempfile.mkdtemp()
        temp_file = os.path.join(temp_dir, 'preview_temp.xlsx')
        shutil.copy2(file_path, temp_file)
        
        # Process the temporary file (hide rows, select sheet, replace keywords)
        # Use the SAME logic as generate to ensure preview = final result
        print(f"[DEBUG] Calling fill_excel_pengalaman with company_data: {company_data}")
        fill_excel_pengalaman(temp_file, company_data, pengalaman_data, {'keywords': keywords})
        
        # Load the PROCESSED workbook
        wb = load_workbook(temp_file, data_only=True)
        
        # Get first visible sheet (after sheet deletion logic)
        sheet = wb.active
        sheet_name = sheet.title
        
        print(f"[DEBUG] After processing - Active sheet: {sheet_name}")
        print(f"[DEBUG] All sheets in workbook: {wb.sheetnames}")
        
        # Convert to HTML table
        html_table = '<table class="table table-bordered table-sm">'
        
        # Track merged cells
        merged_ranges = list(sheet.merged_cells.ranges)
        processed_cells = set()
        
        # Iterate through rows
        max_row = sheet.max_row
        max_col = sheet.max_column
        
        for row_idx in range(1, max_row + 1):
            # Skip hidden rows
            if sheet.row_dimensions[row_idx].hidden:
                continue
                
            html_table += '<tr>'
            
            for col_idx in range(1, max_col + 1):
                # Skip hidden columns
                col_letter = get_column_letter(col_idx)
                if sheet.column_dimensions[col_letter].hidden:
                    continue
                
                cell = sheet.cell(row_idx, col_idx)
                cell_coord = f"{col_letter}{row_idx}"
                
                # Skip if already processed (part of merged cell)
                if cell_coord in processed_cells:
                    continue
                
                # Check if cell is merged
                colspan = 1
                rowspan = 1
                is_merged = False
                
                for merged_range in merged_ranges:
                    if cell.coordinate in merged_range:
                        is_merged = True
                        # Calculate colspan and rowspan
                        min_col, min_row, max_col_merge, max_row_merge = merged_range.bounds
                        colspan = max_col_merge - min_col + 1
                        rowspan = max_row_merge - min_row + 1
                        
                        # Mark all cells in merged range as processed
                        for r in range(min_row, max_row_merge + 1):
                            for c in range(min_col, max_col_merge + 1):
                                processed_cells.add(f"{get_column_letter(c)}{r}")
                        break
                
                # Get cell value
                cell_value = cell.value if cell.value is not None else ''
                
                # Replace keywords
                if isinstance(cell_value, str):
                    for key, value in keywords.items():
                        placeholder = '{' + key + '}'
                        cell_value = cell_value.replace(placeholder, str(value) if value else '')
                
                # Build cell HTML with styling
                style_parts = []
                
                # Text alignment
                if cell.alignment.horizontal:
                    style_parts.append(f"text-align: {cell.alignment.horizontal}")
                if cell.alignment.vertical:
                    style_parts.append(f"vertical-align: {cell.alignment.vertical}")
                
                # Font styling
                if cell.font:
                    if cell.font.bold:
                        style_parts.append("font-weight: bold")
                    if cell.font.italic:
                        style_parts.append("font-style: italic")
                    if cell.font.size:
                        style_parts.append(f"font-size: {cell.font.size}pt")
                
                # Background color
                if cell.fill and cell.fill.start_color:
                    try:
                        rgb = cell.fill.start_color.rgb
                        if isinstance(rgb, str) and rgb != '00000000':  # Not default
                            # Remove alpha channel if present (first 2 chars)
                            if len(rgb) == 8:
                                rgb = rgb[2:]
                            style_parts.append(f"background-color: #{rgb}")
                    except (AttributeError, TypeError):
                        pass  # Skip if color format is not supported
                
                # Border
                style_parts.append("border: 1px solid #dee2e6")
                style_parts.append("padding: 5px")
                
                style_attr = f' style="{"; ".join(style_parts)}"' if style_parts else ''
                colspan_attr = f' colspan="{colspan}"' if colspan > 1 else ''
                rowspan_attr = f' rowspan="{rowspan}"' if rowspan > 1 else ''
                
                html_table += f'<td{colspan_attr}{rowspan_attr}{style_attr}>{cell_value}</td>'
            
            html_table += '</tr>'
        
        html_table += '</table>'
        
        # Add styling
        styled_html = f"""
        <style>
            .excel-preview {{
                font-family: 'Calibri', Arial, sans-serif;
                font-size: 11pt;
                overflow-x: auto;
            }}
            .excel-preview table {{
                border-collapse: collapse;
                width: 100%;
                margin: 0;
            }}
            .excel-preview td {{
                white-space: pre-wrap;
                word-wrap: break-word;
            }}
        </style>
        <div class="excel-preview">
            <div class="alert alert-info mb-2">
                <i class="fas fa-file-excel me-2"></i>Sheet: <strong>{sheet_name}</strong>
            </div>
            {html_table}
        </div>
        """
        
        # Check for remaining placeholders
        remaining_placeholders = []
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    placeholders = re.findall(r'\{([a-zA-Z0-9_\-]+)\}', cell.value)
                    remaining_placeholders.extend(placeholders)
        
        wb.close()
        
        # Cleanup temp file
        try:
            shutil.rmtree(temp_dir)
        except:
            pass
        
        print(f"[DEBUG] Excel preview generated successfully")
        print(f"[DEBUG] Remaining placeholders: {list(set(remaining_placeholders))}")
        
        return jsonify({
            'success': True,
            'html': styled_html,
            'sheet_name': sheet_name,
            'remaining_placeholders': list(set(remaining_placeholders)),
            'file_path': file_path
        })
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"[ERROR] Excel preview failed:")
        print(error_trace)
        return jsonify({
            'success': False,
            'error': f'Error generating preview: {str(e)}',
            'traceback': error_trace
        }), 500
    except Exception as e:
        print(f"[ERROR] Preview Excel error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Terjadi kesalahan: {str(e)}'
        }), 500

@app.route('/health')
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'version': '2.0.0-keywords'
    })

# Register SPSE crawler endpoint
create_spse_endpoint(app)

if __name__ == '__main__':
    # Create required directories
    os.makedirs('templates', exist_ok=True)
    os.makedirs(PROCESSED_FILES_DIR, exist_ok=True)
    
    app.run(debug=True, host='localhost', port=5001)
