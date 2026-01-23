import re
import logging
from typing import Dict, Any, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .utils import terbilang

logger = logging.getLogger(__name__)


def process_hyperlinks_in_paragraph(paragraph, keywords: Dict[str, Any], keyword_details: Optional[Dict[str, Any]] = None) -> int:
    if keyword_details is None:
        keyword_details = {}

    total_replacements = 0
    para_xml = paragraph._element
    hyperlinks = para_xml.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
    if not hyperlinks:
        return 0

    pattern = re.compile(r'\{([a-zA-Z0-9_\-]+)\}')

    for hyperlink in hyperlinks:
        text_elements = hyperlink.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        full_text = "".join([t.text for t in text_elements if t.text])
        matches = list(pattern.finditer(full_text))
        if not matches:
            continue

        for match in matches:
            keyword = match.group(1)
            matched_key = None
            for k in keywords.keys():
                if k.lower() == keyword.lower():
                    matched_key = k
                    break

            if matched_key:
                value_to_replace = str(keywords[matched_key])
                if text_elements:
                    text_elements[0].text = value_to_replace
                    for t in text_elements[1:]:
                        t.text = ""

                    r_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if r_id:
                        try:
                            doc_part = paragraph.part
                            if hasattr(doc_part, 'rels') and r_id in doc_part.rels:
                                rel = doc_part.rels[r_id]
                                if hasattr(rel, 'target_ref'):
                                    current_target = rel.target_ref
                                    if current_target:
                                        if current_target.startswith('mailto:'):
                                            new_target = f"mailto:{value_to_replace}"
                                            rel._target = new_target
                                            logger.debug(f"[HYPERLINK] Updated mailto: {current_target} -> {new_target}")
                                        elif '{' in current_target:
                                            updated_url = current_target
                                            for kw, val in keywords.items():
                                                updated_url = updated_url.replace(f'{{{kw}}}', str(val))
                                            rel._target = updated_url
                                            logger.debug(f"[HYPERLINK] Updated URL: {current_target} -> {updated_url}")
                        except Exception as e:
                            logger.warning(f"Could not update link address for {matched_key}: {e}")

                    if matched_key in keyword_details:
                        keyword_details[matched_key]['count'] += 1
                    else:
                        keyword_details[matched_key] = {'value': value_to_replace, 'count': 1}

                    total_replacements += 1

    return total_replacements


def process_paragraph_keywords(paragraph, keywords: Dict[str, Any], keyword_details: Optional[Dict[str, Any]] = None) -> int:
    if keyword_details is None:
        keyword_details = {}

    hyperlink_replacements = process_hyperlinks_in_paragraph(paragraph, keywords, keyword_details)
    if not paragraph.runs:
        return hyperlink_replacements

    full_text = "".join(run.text for run in paragraph.runs)
    if '{' not in full_text:
        return hyperlink_replacements

    pattern = re.compile(r'\{([a-zA-Z0-9_\-]+)\}')
    matches = list(pattern.finditer(full_text))
    if not matches:
        return hyperlink_replacements

    total_replacements_in_paragraph = 0
    for match in reversed(matches):
        keyword = match.group(1)
        matched_key = None
        for k in keywords.keys():
            if k.lower() == keyword.lower():
                matched_key = k
                break

        if matched_key:
            value_to_replace = str(keywords[matched_key])
            start, end = match.span()
            affected_runs = []
            current_pos = 0
            for i, run in enumerate(paragraph.runs):
                run_len = len(run.text)
                if current_pos < end and current_pos + run_len > start:
                    affected_runs.append((i, run))
                current_pos += run_len

            if not affected_runs:
                continue

            first_run_index, first_run = affected_runs[0]
            if matched_key in keyword_details:
                keyword_details[matched_key]['count'] += 1
            else:
                keyword_details[matched_key] = {'value': value_to_replace, 'count': 1}

            total_replacements_in_paragraph += 1

            run_start_pos = 0
            first_run_original_text = ""
            for i, run in enumerate(paragraph.runs):
                if i == first_run_index:
                    first_run_original_text = run.text
                    break
                run_start_pos += len(run.text)

            prefix_len = start - run_start_pos
            prefix = first_run_original_text[:prefix_len] if prefix_len > 0 else ""

            last_run_index, last_run = affected_runs[-1]
            run_end_pos = 0
            for i, run in enumerate(paragraph.runs):
                run_end_pos += len(run.text)
                if i == last_run_index:
                    break

            suffix_len = run_end_pos - end
            suffix = last_run.text[-suffix_len:] if suffix_len > 0 else ""

            first_run.text = prefix + value_to_replace + suffix
            for i, run in affected_runs[1:]:
                run.text = ""

            full_text = "".join(run.text for run in paragraph.runs)

    return total_replacements_in_paragraph + hyperlink_replacements


def process_location_expansion(doc, keywords):
    """
    Expand {lokasi_pekerjaan} placeholder in tables into multiple rows 
    (Provinsi, Kabupaten, Kecamatan) based on list_lokasi_pekerjaan.
    Label is placed in the placeholder column, Value in the next column.
    Font is set to Arial 9pt.
    """
    if 'list_lokasi_pekerjaan' not in keywords or not keywords['list_lokasi_pekerjaan']:
        return 0
        
    locations = keywords['list_lokasi_pekerjaan']
    total_expanded = 0
    
    for table in doc.tables:
        target_row_idx = -1
        target_col_idx = -1
        
        # Find the row with {lokasi_pekerjaan}
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if '{lokasi_pekerjaan}' in cell.text:
                    target_row_idx = r_idx
                    target_col_idx = c_idx
                    break
            if target_row_idx != -1:
                break
        
        if target_row_idx != -1:
            target_row = table.rows[target_row_idx]
            last_row = target_row
            
            for loc in locations:
                # Define the 3 rows to add
                rows_data = [
                    ("Provinsi", loc.get('provinsi', '')),
                    ("Kabupaten/Kota", loc.get('kabupaten', '')),
                    ("Kecamatan", loc.get('kecamatan', ''))
                ]
                
                for label, value in rows_data:
                    new_row = table.add_row()
                    
                    # Logic: Label at target_col_idx, Value at target_col_idx + 1
                    # Ensure we have enough columns
                    if target_col_idx + 1 < len(new_row.cells):
                        # Label Cell
                        cell_label = new_row.cells[target_col_idx]
                        cell_label.text = label
                        for p in cell_label.paragraphs:
                            p.paragraph_format.left_indent = Cm(0.6)  # Indent 0.5 cm
                            for run in p.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(9)
                        
                        # Value Cell
                        cell_value = new_row.cells[target_col_idx + 1]
                        cell_value.text = value
                        for p in cell_value.paragraphs:
                            for run in p.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(9)
                                
                    elif target_col_idx < len(new_row.cells):
                        # Fallback if next column doesn't exist: put everything in one cell
                        cell = new_row.cells[target_col_idx]
                        cell.text = f"{label}: {value}"
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(9)
                    
                    # Merge columns 4 and 5 (indices 3 and 4) if they exist
                    if len(new_row.cells) >= 5:
                        new_row.cells[3].merge(new_row.cells[4])
                        
                    # Fill remaining columns with "-" centered (after the value column)
                    start_fill_idx = target_col_idx + 2
                    for i in range(start_fill_idx, len(new_row.cells)):
                         cell = new_row.cells[i]
                         cell.text = "-"
                         for p in cell.paragraphs:
                             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                             for run in p.runs:
                                 run.font.name = 'Arial'
                                 run.font.size = Pt(9)

                    # Move new_row to correct position
                    table._element.remove(new_row._element)
                    last_row._element.addnext(new_row._element)
                    last_row = new_row
            
            # Remove the original placeholder row
            table._element.remove(target_row._element)
            total_expanded += 1
            
    return total_expanded


def process_lingkup_expansion(doc, keywords):
    """
    Expand {lingkup_pekerjaan} placeholder in tables into multiple rows 
    based on list_lingkup.
    Data is placed in the placeholder column.
    Font is set to Arial 9pt.
    """
    if 'list_lingkup' not in keywords or not keywords['list_lingkup']:
        return 0
        
    lingkup_list = keywords['list_lingkup']
    total_expanded = 0
    
    for table in doc.tables:
        target_row_idx = -1
        target_col_idx = -1
        
        # Find the row with {lingkup_pekerjaan}
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                if '{lingkup_pekerjaan}' in cell.text:
                    target_row_idx = r_idx
                    target_col_idx = c_idx
                    break
            if target_row_idx != -1:
                break
        
        if target_row_idx != -1:
            target_row = table.rows[target_row_idx]
            last_row = target_row
            
            for item in lingkup_list:
                new_row = table.add_row()
                
                # Logic: Value at target_col_idx
                if target_col_idx < len(new_row.cells):
                    cell_value = new_row.cells[target_col_idx]
                    cell_value.text = item
                    for p in cell_value.paragraphs:
                        for run in p.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
                
                # Merge columns 4 and 5 (indices 3 and 4) if they exist
                if len(new_row.cells) >= 5:
                    new_row.cells[3].merge(new_row.cells[4])
                
                # Fill remaining columns with "-" centered (after the value column)
                start_fill_idx = target_col_idx + 1
                for i in range(start_fill_idx, len(new_row.cells)):
                     cell = new_row.cells[i]
                     cell.text = "-"
                     for p in cell.paragraphs:
                         p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                         for run in p.runs:
                             run.font.name = 'Arial'
                             run.font.size = Pt(9)
                
                # Move new_row to correct position
                table._element.remove(new_row._element)
                last_row._element.addnext(new_row._element)
                last_row = new_row
            
            # Remove the original placeholder row
            table._element.remove(target_row._element)
            total_expanded += 1
            
    return total_expanded


def process_docx_comprehensive(file_path, keywords, output_path):
    """Process DOCX file with comprehensive keyword replacement and row deletion support"""
    try:
        doc = Document(file_path)
        
        # Pre-process dynamic location tables
        process_location_expansion(doc, keywords)
        
        # Pre-process dynamic lingkup tables
        process_lingkup_expansion(doc, keywords)
        
        total_replacements = 0
        log_entries = []
        keyword_details = {}  # Track details for each keyword
        
        # Process main document paragraphs
        for paragraph in doc.paragraphs:
            before_text = paragraph.text
            replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
            after_text = paragraph.text
            
            if replacements > 0:
                total_replacements += replacements
                log_entries.append({
                    'context': 'paragraph',
                    'before': before_text,
                    'after': after_text,
                    'replacements': replacements
                })
        
        # Process tables with deletion support
        for table in doc.tables:
            rows_to_delete = []
            
            for row_idx, row in enumerate(table.rows):
                row_should_delete = False
                
                # Check each cell for "delete" keyword or deleted document references
                for cell in row.cells:
                    cell_text = cell.text.lower().strip()
                    
                    # Check if cell contains "delete"
                    if cell_text == 'delete':
                        row_should_delete = True
                        break
                    
                    if row_should_delete:
                        break
                    
                    # Process keywords in cell paragraphs
                    for paragraph in cell.paragraphs:
                        before_text = paragraph.text
                        replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                        after_text = paragraph.text
                        
                        if replacements > 0:
                            total_replacements += replacements
                            log_entries.append({
                                'context': 'table_cell',
                                'before': before_text,
                                'after': after_text,
                                'replacements': replacements
                            })
                
                if row_should_delete:
                    rows_to_delete.append(row_idx)
            
            # Delete rows in reverse order to maintain correct indices
            for row_idx in reversed(rows_to_delete):
                try:
                    table._element.remove(table.rows[row_idx]._element)
                    log_entries.append({
                        'context': 'table_row_deleted',
                        'before': f'Row {row_idx + 1} in table',
                        'after': 'DELETED',
                        'replacements': 0
                    })
                except:
                    pass
        
        # Process headers and footers
        for section in doc.sections:
            # Headers
            if section.header:
                for paragraph in section.header.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({
                            'context': 'header',
                            'before': before_text,
                            'after': after_text,
                            'replacements': replacements
                        })
                # Header tables
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                before_text = paragraph.text
                                replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                                after_text = paragraph.text
                                if replacements > 0:
                                    total_replacements += replacements
                                    log_entries.append({
                                        'context': 'header_table_cell',
                                        'before': before_text,
                                        'after': after_text,
                                        'replacements': replacements
                                    })
            
            # First Page Header
            if section.first_page_header:
                for paragraph in section.first_page_header.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({
                            'context': 'first_page_header',
                            'before': before_text,
                            'after': after_text,
                            'replacements': replacements
                        })
                # First page header tables
                for table in section.first_page_header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                before_text = paragraph.text
                                replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                                after_text = paragraph.text
                                if replacements > 0:
                                    total_replacements += replacements
                                    log_entries.append({
                                        'context': 'first_page_header_table_cell',
                                        'before': before_text,
                                        'after': after_text,
                                        'replacements': replacements
                                    })
            
            # Even Page Header
            if section.even_page_header:
                for paragraph in section.even_page_header.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({
                            'context': 'even_page_header',
                            'before': before_text,
                            'after': after_text,
                            'replacements': replacements
                        })
                # Even page header tables
                for table in section.even_page_header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                before_text = paragraph.text
                                replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                                after_text = paragraph.text
                                if replacements > 0:
                                    total_replacements += replacements
                                    log_entries.append({
                                        'context': 'even_page_header_table_cell',
                                        'before': before_text,
                                        'after': after_text,
                                        'replacements': replacements
                                    })
            
            # Footers
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({
                            'context': 'footer',
                            'before': before_text,
                            'after': after_text,
                            'replacements': replacements
                        })
                # Footer tables
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                before_text = paragraph.text
                                replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                                after_text = paragraph.text
                                if replacements > 0:
                                    total_replacements += replacements
                                    log_entries.append({
                                        'context': 'footer_table_cell',
                                        'before': before_text,
                                        'after': after_text,
                                        'replacements': replacements
                                    })
        
        # Save the processed document
        doc.save(output_path)
        
        return True, {
            'total_replacements': total_replacements,
            'log_entries': log_entries,
            'keyword_details': keyword_details
        }
        
    except Exception as e:
        return False, {'error': str(e)}


def process_docx_keywords(file_path, keywords, output_path):
    """Process DOCX file and replace all keywords (legacy function)"""
    try:
        doc = Document(file_path)
        total_replacements = 0
        log_entries = []
        keyword_details = {}

        for paragraph in doc.paragraphs:
            before_text = paragraph.text
            replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
            after_text = paragraph.text
            if replacements > 0:
                total_replacements += replacements
                log_entries.append({'context': 'paragraph', 'before': before_text, 'after': after_text, 'replacements': replacements})

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        before_text = paragraph.text
                        replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                        after_text = paragraph.text
                        if replacements > 0:
                            total_replacements += replacements
                            log_entries.append({'context': 'table_cell', 'before': before_text, 'after': after_text, 'replacements': replacements})

        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({'context': 'header', 'before': before_text, 'after': after_text, 'replacements': replacements})
            if section.first_page_header:
                for paragraph in section.first_page_header.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({'context': 'first_page_header', 'before': before_text, 'after': after_text, 'replacements': replacements})
            if section.even_page_header:
                for paragraph in section.even_page_header.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({'context': 'even_page_header', 'before': before_text, 'after': after_text, 'replacements': replacements})
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    before_text = paragraph.text
                    replacements = process_paragraph_keywords(paragraph, keywords, keyword_details)
                    after_text = paragraph.text
                    if replacements > 0:
                        total_replacements += replacements
                        log_entries.append({'context': 'footer', 'before': before_text, 'after': after_text, 'replacements': replacements})

        doc.save(output_path)
        return True, {'total_replacements': total_replacements, 'log_entries': log_entries, 'keyword_details': keyword_details}
    except Exception as e:
        return False, {'error': str(e)}
